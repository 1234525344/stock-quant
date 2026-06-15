#!/usr/bin/env python
"""Rebuild 谭欣宇实习总结3.0 per teacher template structure, preserving cover."""
import sys, os, copy, re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Mm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

SRC = r'C:/Users/lb/Desktop/00232/谭欣宇实习总结3.0.docx'
DST = r'C:/Users/lb/Desktop/00232/谭欣宇实习总结3.0_重构.docx'
try: os.remove(DST)
except: pass

FANG = '仿宋_GB2312'
HEI = '黑体'

# ==================== BUILD ====================
doc = Document()

# ---- FUNCTIONS ----
def mk_font(run, fn, sz, bold=False):
    run.font.name = fn; run.font.size = sz; run.bold = bold
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None: rPr = OxmlElement('w:rPr'); run._r.insert(0, rPr)
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:eastAsia'), fn); rF.set(qn('w:ascii'), fn); rF.set(qn('w:hAnsi'), fn)
    rPr.insert(0, rF)

def P(text='', fn=FANG, sz=Pt(14), bold=False, align=None, indent=0, sb=0, sa=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.line_spacing = Pt(28); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    if indent: pf.first_line_indent = Pt(indent)
    if align == 'C': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        r = p.add_run(text); mk_font(r, fn, sz, bold)
    return p

def T(text): return P(text, HEI, Pt(22), True, 'C', 0, 6, 10)
def H(text): return P(text, HEI, Pt(14), True, 0, 10, 4)
def B(text): return P(text, FANG, Pt(14), False, 0, 28)
def B0(text): return P(text, FANG, Pt(14), False, 0)
def blank(): return P('', 0)

def add_toc_item(p, name, page):
    """Set paragraph with right-aligned dot-leader tab + page number."""
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None: pPr = OxmlElement('w:pPr'); p._element.insert(0, pPr)
    old = pPr.find(qn('w:tabs'))
    if old is not None: pPr.remove(old)
    tabs = OxmlElement('w:tabs')
    tb = OxmlElement('w:tab'); tb.set(qn('w:val'), 'right'); tb.set(qn('w:pos'), '8844'); tb.set(qn('w:leader'), 'dot')
    tabs.append(tb); pPr.append(tabs)
    for r in p.runs: r.text = ''
    if p.runs: p.runs[0].text = f'{name}\t{page}'
    else:
        r = p.add_run(f'{name}\t{page}'); mk_font(r, FANG, Pt(14))

def toc_hyperlink(p, name, bm_name, page=''):
    """Add clickable TOC entry that jumps to bookmark."""
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None: pPr = OxmlElement('w:pPr'); p._element.insert(0, pPr)
    old = pPr.find(qn('w:tabs'))
    if old is not None: pPr.remove(old)
    tabs = OxmlElement('w:tabs')
    tb = OxmlElement('w:tab'); tb.set(qn('w:val'), 'right'); tb.set(qn('w:pos'), '8844'); tb.set(qn('w:leader'), 'dot')
    tabs.append(tb); pPr.append(tabs)
    # Clear old
    for r in list(p.runs): r._r.getparent().remove(r._r)
    # Hyperlink run
    hl = OxmlElement('w:hyperlink'); hl.set(qn('w:anchor'), bm_name); hl.set(qn('w:history'), '1')
    rh = OxmlElement('w:r')
    rPr2 = OxmlElement('w:rPr')
    rF2 = OxmlElement('w:rFonts'); rF2.set(qn('w:eastAsia'), FANG); rF2.set(qn('w:ascii'), FANG); rF2.set(qn('w:hAnsi'), FANG)
    rPr2.append(rF2); sz2 = OxmlElement('w:sz'); sz2.set(qn('w:val'), '28'); rPr2.append(sz2)
    rh.append(rPr2)
    th = OxmlElement('w:t'); th.set(qn('xml:space'), 'preserve'); th.text = name
    rh.append(th); hl.append(rh); p._element.append(hl)
    # Page number run
    if page:
        tr = OxmlElement('w:r')
        trPr = OxmlElement('w:rPr')
        trF = OxmlElement('w:rFonts'); trF.set(qn('w:eastAsia'), FANG); trF.set(qn('w:ascii'), FANG); trF.set(qn('w:hAnsi'), FANG)
        trPr.append(trF); tsz = OxmlElement('w:sz'); tsz.set(qn('w:val'), '28'); trPr.append(tsz)
        tr.append(trPr)
        t2 = OxmlElement('w:t'); t2.set(qn('xml:space'), 'preserve'); t2.text = '\t' + page
        tr.append(t2); p._element.append(tr)

def add_bm(para, bm_name, bm_id):
    """Add bookmark to paragraph."""
    bs = OxmlElement('w:bookmarkStart'); bs.set(qn('w:id'), str(bm_id)); bs.set(qn('w:name'), bm_name)
    be = OxmlElement('w:bookmarkEnd'); be.set(qn('w:id'), str(bm_id))
    para._element.insert(0, bs); para._element.append(be)

# ===== SECTION 1: COVER =====
s0 = doc.sections[0]
for a, v in [('page_width', Mm(210)), ('page_height', Mm(297)),
    ('top_margin', Mm(35)), ('bottom_margin', Mm(32)), ('left_margin', Mm(28)), ('right_margin', Mm(26))]:
    setattr(s0, a, v)

# ---- Build cover (clone from original) ----
# Table 0: School name
t0 = doc.add_table(rows=1, cols=2)
t0.cell(0,1).text = '江西职业技术大学'
for p in t0.cell(0,1).paragraphs:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(22); r.font.name = '黑体'; r.font.bold = True
blank(); blank()
P('JIANGXI POLYTECHNIC UNIVERSITY', HEI, Pt(15), True, 'C')
P('船舶工程学院岗位实习报告', HEI, Pt(24), True, 'C', 0, 4, 40)

# Table 1: Student info
blank(); blank(); blank()
info = [
    ('专业：', '供热通风与空调工程技术'),
    ('班级：', '空调2301'),
    ('学生姓名：', '谭欣宇'),
    ('学号：', '232030311'),
    ('企业指导教师：', ''),
    ('学校指导教师：', '张伟明'),
    ('实习单位：', '进贤申通快递'),
    ('实习形式：', '□统一安排实习     ☑自主实习'),
]
t1 = doc.add_table(rows=len(info), cols=2)
for i, (k, v) in enumerate(info):
    t1.cell(i,0).text = k
    t1.cell(i,1).text = v
    for p in t1.cell(i,0).paragraphs:
        for r in p.runs: r.font.size = Pt(14); r.font.name = FANG
    for p in t1.cell(i,1).paragraphs:
        for r in p.runs: r.font.size = Pt(14); r.font.name = FANG

blank(); blank(); blank()
# Table 2: Evaluation
t2 = doc.add_table(rows=2, cols=2)
t2.cell(0,0).text = '校内指导教师评价'
t2.cell(0,1).text = '教师签名：'
t2.cell(1,0).text = '实习成绩'
for row in t2.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(14); r.font.name = FANG

blank(); blank()
P('二〇二六年六月', FANG, Pt(14), False, 'C', 0, 30)

# ===== SECTION 2: BODY =====
doc.add_section()
s1 = doc.sections[1]
for a, v in [('page_width', Mm(210)), ('page_height', Mm(297)),
    ('top_margin', Mm(35)), ('bottom_margin', Mm(32)), ('left_margin', Mm(28)), ('right_margin', Mm(26))]:
    setattr(s1, a, v)

# Page number
pg = OxmlElement('w:pgNumType'); pg.set(qn('w:start'), '1'); s1._sectPr.append(pg)
ft = s1.footer; ft.is_linked_to_previous = False
fp = ft.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(); fr.font.name = FANG; fr.font.size = Pt(14)
rPr = fr._r.find(qn('w:rPr'))
if rPr is None: rPr = OxmlElement('w:rPr'); fr._r.insert(0, rPr)
rF = OxmlElement('w:rFonts')
rF.set(qn('w:eastAsia'), FANG); rF.set(qn('w:ascii'), FANG); rF.set(qn('w:hAnsi'), FANG)
rPr.insert(0, rF)
for t, v in [('b','begin'),('t',' PAGE '),('e','end')]:
    el = OxmlElement('w:fldChar') if t != 'e' and t != 'b' else OxmlElement('w:fldChar')
    if t == 't': el = OxmlElement('w:instrText'); el.set(qn('xml:space'),'preserve'); el.text = v
    else: el.set(qn('w:fldCharType'), v)
    fr._r.append(el)

# ===== TOC =====
T('目  录')
bm_counter = [100]  # mutable counter

def toc_entry(name, page):
    p = P(indent=0)
    bm = 'bm' + str(bm_counter[0]); bm_counter[0] += 1
    add_toc_item(p, name, page)
    return p, bm

def toc_section(name, page):
    p = P(name, FANG, Pt(14))
    bm = 'bm' + str(bm_counter[0]); bm_counter[0] += 1
    toc_hyperlink(p, name, bm, page)
    return p, bm

def toc_chapter(name, page):
    p = P(name, HEI, Pt(22), True)
    bm = 'bm' + str(bm_counter[0]); bm_counter[0] += 1
    toc_hyperlink(p, name, bm, page)
    return p, bm

# Build TOC
toc_entries = [
    ('chapter', '一、综述', '1'),
    ('sub', '  （一）实习单位基本情况', '1'),
    ('sub', '  （二）实习岗位描述', '2'),
    ('chapter', '二、主体', '3'),
    ('sub', '  （一）职业素养与社会能力提升', '3'),
    ('sub', '  （二）实习内容、收获与技能掌握', '4'),
    ('sub', '  （三）发现、分析与解决实际问题的能力简述', '6'),
    ('sub', '  （四）专业技能与岗位要求的匹配度与差距', '7'),
    ('sub', '  （五）对专业课程与实习安排的建议', '8'),
    ('sub', '  （六）今后努力方向', '9'),
    ('chapter', '三、总结', '10'),
]

toc_bms = {}  # name -> (para_element, bm_name)
for typ, name, page in toc_entries:
    if typ == 'chapter':
        p, bm = toc_chapter(name, page)
    else:
        p, bm = toc_section(name, page)
    toc_bms[name] = (p, bm)

doc.add_page_break()

# ===== 一、综述 =====
p, bm = toc_bms['一、综述']; T('一、综述'); add_bm(doc.paragraphs[-1], bm, bm_counter[0]-1)

p2, bm2 = toc_bms['  （一）实习单位基本情况']; H('（一）实习单位基本情况'); add_bm(doc.paragraphs[-1], bm2, bm_counter[0]-1)
B('本次实习单位为申通快递进贤分公司，位于江西省南昌市进贤县。进贤县作为连接南昌市区与周边县市的重要交通枢纽，其物流快递业务量近年来随着电商下沉市场的蓬勃发展而日益增长。申通快递作为国内最早成立的民营快递企业之一，成立于1993年，拥有覆盖全国的庞大服务网络和深厚的品牌积淀。进贤分公司作为申通快递网络中的末端网点，承担着全县范围内快件的揽收、中转与派送重任。虽然进贤分公司属于基层网点，规模无法与省级转运中心相比，但其运作模式高度浓缩了快递行业的核心要素。网点内部架构清晰，设有操作部、客服部、市场部及车队。操作部负责快件的进出港处理，是网点的核心部门；客服部负责处理查单、投诉与理赔；市场部则负责拓展电商客户与散单业务。公司依托申通总部强大的信息技术支持，实现了快件流转的全程可视化。在进贤当地，申通快递凭借稳定的时效和广泛的服务覆盖，与多家本地特色产品商家建立了长期合作关系。网点核心团队经验丰富，对于应对"双十一""六一八"等物流高峰有着成熟的预案和应对机制。')

p3, bm3 = toc_bms['  （二）实习岗位描述']; H('（二）实习岗位描述'); add_bm(doc.paragraphs[-1], bm3, bm_counter[0]-1)
B('本次顶岗实习中，我主要担任快递操作员兼客服助理的复合型基础岗位。该岗位要求实习生既要具备体力和耐力，又要具备细心和沟通能力。具体职责包括以下两方面：')
B('在操作端方面，负责进出港快件的卸车、扫描、分拣与装袋；熟练使用巴枪进行快件信息的录入与更新；协助处理破损件、无头件等异常快件；参与网点车辆的装载与卸载工作。在客服端方面，协助客服主管处理客户查件请求；通过电话或微信与客户沟通派送时间与地点；处理简单的客户投诉与咨询；协助进行电商大客户的对账单整理。')
B('该岗位要求我必须熟悉申通快递的内部操作系统，了解《邮政法》及快递服务国家标准中关于时限、赔偿的相关规定，同时还需要具备良好的身体素质和抗压能力，以适应快递行业早出晚归、节假日无休的工作常态。正是在这样一个综合性的岗位上，我从最初的手忙脚乱逐步成长为能够独立处理日常操作和客户问题的准职业人。')

# ===== 二、主体 =====
p4, bm4 = toc_bms['二、主体']; T('二、主体'); add_bm(doc.paragraphs[-1], bm4, bm_counter[0]-1)

# -- (一) 职业素养 --
p5, bm5 = toc_bms['  （一）职业素养与社会能力提升']; H('（一）职业素养与社会能力提升'); add_bm(doc.paragraphs[-1], bm5, bm_counter[0]-1)

H('1. 规章制度与劳动纪律')
B('实习期间，我严格遵守申通快递进贤分公司的各项规章制度和操作规范。快递行业是一个高度规范化的行业，从巴枪扫描的操作流程到快件交接的签收制度，每一个环节都有明确的标准和要求。在出勤方面，我坚持每天早晨六点前到岗参加早班分拣，从未出现迟到、早退或旷工现象。在操作规范方面，我严格执行"先扫描、后分拣"的工作流程，杜绝先分后扫或漏扫现象，确保每一个包裹的流转信息完整可追溯。在处理异常快件时，严格按照公司规定的"拍照—上报—登记—隔离"四步流程操作，绝不自作主张。')

H('2. 职业道德')
B('快递行业是一个与客户直接接触的服务行业，诚信和责任心是最基本的职业道德要求。实习初期，我曾因追求扫描效率而在未实际派送的情况下提前点击"签收"，导致客户投诉和网点罚款。这次教训让我深刻认识到"虚假签收"是快递行业的红线，物流信息的真实性是建立客户信任的基础。从此以后，我始终坚持实事求是的工作原则，每一条操作记录都确保与实际物流状态一致。在经手客户包裹时，我始终像对待自己的物品一样小心谨慎，用实际行动践行快递行业的职业操守。')

H('3. 社会适应能力与存在不足')
B('进入进贤申通快递的最初几周，我从相对宽松的校园环境切换到高强度、快节奏的快递操作现场，经历了较大的身心冲击。每天搬运上千个包裹导致腰背酸痛、双手磨出水泡，甚至一度产生过放弃的念头。但我调整心态，将其视为对自己意志力的磨练，同时观察老员工的操作技巧，学会了利用身体重心和腿部力量搬运重物，逐渐适应了这种高强度的工作节奏。在日常工作中，我注重与快递员、客服同事、主管的沟通配合，学会了在忙碌中清晰传递信息、在分歧中主动协调。')
B('不足之处在于，实习初期在客服沟通方面表现欠佳。在处理一位客户关于包裹延误的投诉时，我因急于解释是天气原因而打断了客户的抱怨，语气显得不耐烦，导致客户升级投诉。主管教导我，客服的第一原则是"共情"，先安抚情绪再解决问题。这次经历让我认识到自己在职场沟通技巧上的稚嫩，还需要在换位思考和情绪管理方面持续提升。')

# -- (二) 实习内容、收获 --
p6, bm6 = toc_bms['  （二）实习内容、收获与技能掌握']; H('（二）实习内容、收获与技能掌握'); add_bm(doc.paragraphs[-1], bm6, bm_counter[0]-1)

H('1. 主要工作内容')
B('（1）早班分拣与出仓操作（6:00—9:00）。每天清晨，来自南昌转运中心的大型货车抵达网点，我的第一项任务是协助卸车。这是一项体力活，需要快速将传送带上的包裹卸下。随后进入核心的分拣环节，根据面单上的地址将包裹快速归类到对应的片区格口（如民和镇片区、温圳片区、李渡片区等），同时通过巴枪扫描每一个包裹，确保系统显示"正在派送"。对易碎品和大件货物进行特殊标记和单独存放。')
B('（2）中班派送协助与客服（9:00—14:00）。分拣完毕后，我转入客服辅助工作。在电脑前监控系统异常数据，处理因地址模糊或电话空号导致无法派送的包裹，尝试通过系统预留的其他联系方式联系客户。同时处理"拦截件"，即客户在发货后突然要求退回的包裹，在系统中进行操作并通知操作部将其拦截。午间时段协助处理上门寄件的客户，指导填写电子面单并进行称重计费。')
B('（3）晚班操作与入库（14:00—18:00）。下午主要处理进港的二次分拣和驿站入库。操作"菜鸟驿站"等系统的入库端口，扫描包裹条码，生成取件码，并将包裹整齐摆放在货架上。对于无法入库的偏远乡镇件，则重新打包，装上当晚发往乡镇代理点的班车。')

H('2. 收获与体会')
B('（1）快递业务流程与信息系统的深度学习。在实习过程中，我重点学习了申通快递的信息化管理系统。巴枪的高级应用方面，我学会了使用巴枪进行问题件上报（包装破损、面单磨损等），拍照上传系统作为理赔依据，使用地图导航功能辅助规划最优派送路线，以及在签收环节规范录入代收人信息避免虚假签收投诉。内部管理系统方面，我深入学习了如何查看网点的各项KPI指标（签收率、及时率、投诉率等），了解了每一个指标波动对网点考核罚款的影响，学会了在遇到暴雨、道路施工等不可抗力时及时在系统中进行延误报备以免除罚款。')
B('（2）B2B与B2C物流模式差异的实践认知。针对进贤当地的电商客户（如医疗器械厂），我学习了如何配置热敏打印机，如何根据客户ERP系统需求调整电子面单模板（添加店铺Logo、备注信息等）。这让我对B2B物流与B2C物流在操作层面的差异有了直观的认识——前者注重批量处理和系统对接，后者注重末端服务和客户体验。')
B('（3）职场思维方式从学生到职业人的转变。在学校学习时，我更关注个体的学习成果。但在快递网点的实际工作中，每一个环节都紧密相连——卸车慢了会影响分拣，分拣错了会导致错发，签收不规范会引发投诉。我学会了从整体流程的角度思考问题，理解了自己岗位在整个运营链条中的位置和作用。')

H('3. 技能掌握程度')
B('通过本次实习，我在以下方面取得了较为扎实的进步：快递操作方面，能够独立完成快件的卸车、分拣、扫描和入库全流程操作；信息系统方面，熟练掌握了巴枪和内部管理系统的常用功能；客服能力方面，掌握了基本的客户查件处理和投诉应对技巧；物流知识方面，深入了解了快递行业从揽收到派送的全链路运作模式。但在数据分析、管理统筹和复杂客户谈判方面仍有较大提升空间。')

# -- (三) 发现、分析与解决问题 --
p7, bm7 = toc_bms['  （三）发现、分析与解决实际问题的能力简述']; H('（三）发现、分析与解决实际问题的能力简述'); add_bm(doc.paragraphs[-1], bm7, bm_counter[0]-1)

B('在实习期间，我在日常工作中遇到了多个实际问题，并尝试进行分析和解决，以下选取两个典型案例加以说明。')

B('案例一：分拣地址库混乱导致错发问题。问题发现：在早班分拣时，我发现经常有包裹因地址信息不完整（只写了村名没有门牌号）或地名相似（如"温圳镇"与"文港镇"在发音上容易混淆）而被错误分配，导致错发率偏高，影响了客户体验和网点运营成本。问题分析：经与快递员和主管交流，我发现根源在于进贤县地形复杂、乡镇村落众多，很多乡村没有标准门牌号，新员工和外地实习生主要依赖面单信息进行判断，缺乏本地地理知识的支撑。解决方案：我利用休息时间，手绘了一份进贤县主要乡镇的派送路线图，标注了容易混淆的村名和标志性建筑作为参照点。同时向负责各片区的老快递员请教，记住了一些"老地名"。通过"烂笔头"和"勤张嘴"，我逐渐建立了自己的"脑中地图"，分拣准确率显著提高。反思：这次经历让我认识到，物流末端配送的难点往往不在于技术，而在于信息的不对称和本地化知识的缺失。解决这类问题不能只靠系统，还需要一线操作人员主动积累经验、建立知识库。')

B('案例二：系统故障时的应急处理。问题发现：在快递高峰期，巴枪偶尔会出现信号不好或系统崩溃的情况，导致数据无法实时上传，现场出现包裹堆积和秩序混乱。问题分析：系统故障在物流高峰期间是常见问题，但关键在于是否有有效的应急预案。我发现网点虽然有口头约定的应急流程，但缺乏书面规范和实操演练，导致每次故障时的处理效率参差不齐。解决方案：在主管的指导下，我参与了应急预案的优化工作：在系统故障时，立即启动手工记录模式，使用纸质表格登记包裹信息，先进行物理分拣，待系统恢复后再集中补录数据。同时，我学会了检查巴枪的网络设置，在4G和Wi-Fi之间切换，确保设备处于最佳工作状态。这些措施使故障期间的包裹处理效率提升了约30%。反思：这次经历让我深刻认识到，物流运营中的风险管理与日常操作同样重要，应急预案不能只停留在口头，需要书面化、标准化，并定期演练。')

# -- (四) 专业技能与岗位匹配度 --
p8, bm8 = toc_bms['  （四）专业技能与岗位要求的匹配度与差距']; H('（四）专业技能与岗位要求的匹配度与差距'); add_bm(doc.paragraphs[-1], bm8, bm_counter[0]-1)

H('1. 适应情况')
B('通过在校期间学习的《物流管理基础》《供应链管理》《运输管理》等专业课程，我具备了物流与供应链的基本理论知识框架，包括快件流转路径优化、库存控制、客户服务管理等。这些理论知识使我在理解申通快递的运营模式和业务流程时有了较好的基础。例如，学习过的物流信息系统知识帮助我较快地掌握了巴枪操作和网点管理系统的使用；学习过的客户服务理论使我在处理客户投诉时能够从服务和营销的视角进行分析。在校期间培养的计算机操作能力和数据处理意识也帮助我快速适应了快递行业信息化管理的工作环境。')

H('2. 存在的差距')
B('（1）物流数据分析能力不足。当前快递行业的竞争已经进入数据驱动阶段，网点管理越来越依赖KPI数据的分析来优化运营。我在数据收集、统计分析和可视化呈现方面的能力还比较薄弱，距离能够利用数据为网点运营决策提供支持的岗位要求还有较大差距。')
B('（2）物流规划与优化能力欠缺。课堂上学习了物流规划和路径优化的理论知识，但在实际操作层面，如何针对进贤县复杂的乡镇路网进行派送路径的科学规划、如何优化网点内部的分拣流程以提高作业效率等方面的能力还需要大量实践积累。')
B('（3）供应链整体视野不够开阔。目前对快递行业的理解主要停留在末端网点的操作层面，对于上游的干线运输、转运中心运作、仓储管理以及供应链金融等更高层次的内容了解不足，全局视野有待拓展。')
B('（4）综合管理能力空白。在人员调度、绩效考核、成本控制、客户关系管理等管理层面的知识和经验基本空白，这对于未来向管理岗位发展是一大短板。')

# -- (五) 建议 --
p9, bm9 = toc_bms['  （五）对专业课程与实习安排的建议']; H('（五）对专业课程与实习安排的建议'); add_bm(doc.paragraphs[-1], bm9, bm_counter[0]-1)

H('1. 课程设置与安排')
B('（1）加强物流信息系统的实训教学。建议在物流管理专业课程中增加快递行业常用信息系统（如巴枪模拟操作、WMS系统、TMS系统等）的实训内容，让学生在校期间就能接触到行业真实使用的信息化工具，缩短入职适应期。')
B('（2）增开物流数据分析课程。随着快递行业数字化转型的深入，数据分析能力已成为物流从业人员的核心技能之一。建议增设物流数据分析、数据可视化等方面的课程或实训项目。')
B('（3）增加电商物流和快递运营的专题内容。快递行业与电子商务深度绑定，建议在课程中增加电商大促物流保障、逆向物流、社区团购物流等时效性较强的专题内容，紧跟行业发展动态。')

H('2. 教师教学')
B('（1）多采用快递企业的实际案例进行教学。快递行业的变化速度快、热点问题多，建议教师在教学中多选取行业最新案例（如双十一物流保障、快递价格战、末端配送模式创新等）进行分析讨论，增强课程的时效性和吸引力。')
B('（2）强调操作规范和安全意识。快递行业涉及大量一线操作，安全意识和规范操作是基础保障。建议在教学过程中加强这方面的实训和考核。')

H('3. 顶岗实习安排')
B('（1）建议在实习前安排一周左右的行业认知培训，帮助学生了解快递行业的基本运作模式、岗位要求和安全规范，减少实习初期的不适应。')
B('（2）建议在实习中期安排一次集中的返校交流或线上交流，帮助学生梳理已掌握的技能、发现存在的不足，并及时调整学习方向。这次实习中我深切体会到，如果没有中期的反思和总结，很容易陷入日复一日的机械操作而忽略了学习和提升。')

# -- (六) 今后努力方向 --
p10, bm10 = toc_bms['  （六）今后努力方向']; H('（六）今后努力方向'); add_bm(doc.paragraphs[-1], bm10, bm_counter[0]-1)

B('基于对自身不足的清醒认知和对物流行业发展方向的把握，我明确了今后需要努力的方向。')
B('第一，深化物流专业知识的学习。系统学习物流规划与设计、供应链管理、物流信息系统等核心专业课程，夯实理论基础，争取考取物流师等相关职业资格证书，为从事物流行业的高层次工作做好准备。')
B('第二，提升数据分析能力。学习Excel高级功能、SQL数据库查询和Python等数据分析工具，掌握物流数据处理和分析的基本方法，能够利用数据为物流运营决策提供支撑，跟上物流行业数字化转型的步伐。')
B('第三，加强现场实践经验的积累。未来继续扎根物流一线，深入了解不同物流模式和场景的实际运作。关注无人机配送、自动化分拣、无人驿站等物流新技术的应用发展，保持对行业前沿的敏感度。')
B('第四，培养综合管理和沟通协调能力。通过参与更多的团队项目和跨部门协作，提升管理统筹和人际沟通的能力，为未来从操作岗位向管理岗位发展做好准备。')
B('第五，保持终身学习的职业态度。物流行业是技术更新和社会变革的前沿领域，我将始终保持开放的学习心态，通过阅读行业报告、参加专业培训、关注物流科技动态等多种途径，持续提升自己的专业素养和综合能力。')

# ===== 三、总结 =====
p11, bm11 = toc_bms['三、总结']; T('三、总结'); add_bm(doc.paragraphs[-1], bm11, bm_counter[0]-1)

B('时光荏苒，在进贤申通快递的顶岗实习已近尾声。回首这段日子，从最初的懵懂无知、手忙脚乱，到现在的熟练从容、独当一面，我收获的不仅仅是物流操作技能，更是一份沉甸甸的人生感悟。')
B('第一，专业知识的落地与升华。在学校里，我们学习过《供应链管理》《物流运筹学》，那些复杂的模型和理论曾让我觉得枯燥。然而，在申通快递的每一个环节，我都看到了这些理论的影子。分拣路径的优化就是运筹学的实际应用；库存的控制就是供应链管理的体现；客户的投诉处理就是服务营销的实战。我深刻体会到，理论是灰色的，而生命之树常青，只有将理论应用到实践中，知识才能转化为力量。物流不仅仅是"送快递"，而是一个集信息流、资金流、物流于一体的复杂生态系统。')
B('第二，职业素养的养成。"细节决定成败，态度决定一切。"这是我在实习中最大的体会。快递行业是一个服务行业，每一个包裹背后都是一个期待的客户。准时、安全、微笑，这些看似简单的要求，做到极致就是专业。我学会了守时，因为班车不等人；我学会了严谨，因为一个数字的错误可能导致包裹发往千里之外；我学会了忍耐，面对客户的误解和指责，用专业和耐心去化解。这些职业素养，将是我未来无论从事何种行业都必须具备的基石。')
B('第三，对社会与生活的深刻认知。在进贤这片土地上，我接触到了形形色色的人——有为了生计奔波的电商小老板，有在大城市打工给家里寄特产的游子，也有因为一个快递延误而大发雷霆的普通市民。我看到了生活的艰辛，也看到了人性的温暖。快递小哥们风里来雨里去的身影，让我明白了劳动的光荣与不易。这段经历让我褪去了大学生的书生气和娇气，多了一份对社会的敬畏和对劳动者的尊重。')
B('第四，对未来的规划与展望。通过这次实习，我明确了未来的职业发展方向。我意识到自己在数据分析和管理统筹方面还有很大的提升空间。未来，我希望能从事物流规划或供应链数据分析相关的工作，利用数字化手段解决物流痛点，提高物流效率。同时，物流技术日新月异，无人机配送、自动化分拣、大数据预测等技术正在改变行业，我必须保持学习的热情，才能不被时代淘汰。')
B('最后，衷心感谢进贤申通快递给我提供了这样一个宝贵的实习平台，感谢主管和同事们对我的包容与教导。这段充满汗水与欢笑的实习经历，将成为我记忆中最宝贵的财富，激励我在未来的人生道路上，脚踏实地，勇往直前。路虽远，行则将至；事虽难，做则必成。')

# ===== SAVE =====
doc.save(DST)

# Count
total = sum(len(p.text.replace(' ','').replace('\n','')) for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            total += len(c.text.replace(' ','').replace('\n',''))
print(f'Saved: {len(doc.paragraphs)}p + {len(doc.tables)}t, ~{total}c')
print(f'Path: {DST}')
print('DONE')
