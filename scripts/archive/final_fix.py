import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

path = r'C:/Users/lb/Desktop/电力系统工程课程设计_完整版.docx'
doc = Document(path)

F = '仿宋_GB2312'; H = '黑体'

# ===== 1. Remove duplicate TOC + duplicate 评分表 (paras 228-248) =====
# Delete from 评分表 title [228] to end
to_del = list(range(228, len(doc.paragraphs)))
for i in reversed(to_del):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
print('1. Removed duplicate TOC + 评分表')

# ===== 2. Insert section break BEFORE body, add page numbers =====
# Body starts at '1. 住宅用电负荷分类及用电特点' 
body_start = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith('1. 住宅用电负荷分类'):
        body_start = i
        break

# Insert section break in the paragraph just before body
prev_para = doc.paragraphs[body_start - 1]
pPr = prev_para._element.find(qn('w:pPr'))
if pPr is None: pPr = OxmlElement('w:pPr'); prev_para._element.insert(0, pPr)

# Add sectPr to create section break
sectPr = OxmlElement('w:sectPr')
# Copy page setup from current section
ps = OxmlElement('w:pgSz'); ps.set(qn('w:w'), '11907'); ps.set(qn('w:h'), '16840')
sectPr.append(ps)
pm = OxmlElement('w:pgMar')
pm.set(qn('w:top'), '992'); pm.set(qn('w:bottom'), '907')
pm.set(qn('w:left'), '794'); pm.set(qn('w:right'), '737')
sectPr.append(pm)
# Page number start at 1 for body section
pn = OxmlElement('w:pgNumType'); pn.set(qn('w:start'), '1')
sectPr.append(pn)
# Footer reference (will be fixed after save)
pPr.append(sectPr)
print('2. Section break inserted before body. Sections:', len(doc.sections) if hasattr(doc,'_sections') else 'N/A')

# ===== 3. Save temp and reopen to add footer =====
import tempfile
fd, tmp = tempfile.mkstemp(suffix='.docx'); os.close(fd)
doc.save(tmp)

# Reopen
doc2 = Document(tmp)
# Body is now section 1 (section 0 = cover+TOC)
if len(doc2.sections) >= 2:
    body_sec = doc2.sections[1]
    # Page number
    for c in list(body_sec._sectPr):
        if c.tag == qn('w:pgNumType'): body_sec._sectPr.remove(c)
    pgn = OxmlElement('w:pgNumType'); pgn.set(qn('w:start'), '1')
    body_sec._sectPr.append(pgn)
    
    # Footer with PAGE field
    body_sec.different_first_page_header_footer = False
    ft = body_sec.footer; ft.is_linked_to_previous = False
    fp = ft.paragraphs[0]
    for r in list(fp.runs): fp._p.remove(r._r)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(); fr.font.name = F; fr.font.size = Pt(14)
    rPr = fr._r.find(qn('w:rPr'))
    if rPr is None: rPr = OxmlElement('w:rPr'); fr._r.insert(0, rPr)
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:eastAsia'), F); rF.set(qn('w:ascii'), F); rF.set(qn('w:hAnsi'), F)
    rPr.insert(0, rF)
    for tp, vl in [('b','begin'),('t',' PAGE '),('e','end')]:
        if tp == 't': el = OxmlElement('w:instrText'); el.set(qn('xml:space'),'preserve'); el.text = vl
        else: el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'), vl)
        fr._r.append(el)
    print('3. Page numbers added to body section (仿宋14pt, start=1)')
else:
    print('3. WARN: only 1 section found!')

# ===== 4. Fix chapter heading formatting =====
# Body chapter titles (1. 2. 3. ... 7.) should be 黑体 22pt bold centered
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    if t.startswith('1. 住宅用电负荷') or t.startswith('2. 电力负荷计算') or \
       t.startswith('3. 家庭照明') or t.startswith('4. 供配电系统') or \
       t.startswith('5. 用电安全') or t.startswith('6. 灯具、电气') or \
       t.startswith('7. 总结'):
        for r in p.runs:
            r.font.name = H; r.font.size = Pt(22); r.font.bold = True
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None: pPr = OxmlElement('w:pPr'); p._element.insert(0, pPr)
        jc = pPr.find(qn('w:jc'))
        if jc is None: jc = OxmlElement('w:jc'); pPr.append(jc)
        jc.set(qn('w:val'), 'center')
        # Remove indent
        ind = pPr.find(qn('w:ind'))
        if ind is not None: pPr.remove(ind)
        print(f'  Fixed chapter [{i}]: {t[:40]}')

# ===== 5. Fix section headers (1.1, 1.2, ...) =====
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    if t.startswith('1.1 ') or t.startswith('1.2 ') or t.startswith('1.3 ') or \
       t.startswith('2.1 ') or t.startswith('2.2 ') or \
       t.startswith('表') and any(c.isdigit() for c in t):
        for r in p.runs:
            if r.font.size is None or (r.font.size or 0)/12700 < 14:
                r.font.name = H; r.font.size = Pt(14); r.font.bold = True

outpath = r'C:/Users/lb/Desktop/电力系统工程课程设计_完整版.docx'
try: os.remove(outpath)
except: pass
doc2.save(outpath)
# Cleanup
os.remove(tmp)

print(f'\nSaved: {outpath}')
print(f'Sections: {len(doc2.sections)}, Paras: {len(doc2.paragraphs)}, Tables: {len(doc2.tables)}')
# Image count
from docx.oxml.ns import qn
imgs = sum(1 for p in doc2.paragraphs for r in p.runs if r._r.findall(qn('w:drawing')))
print(f'Images: {imgs}')
