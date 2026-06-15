import json, sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.17); s.right_margin = Cm(3.17)

SONGTI = '宋体'; HEI = '黑体'; KAITI = '楷体'
LINE_28 = Pt(28); LINE_36 = Pt(36)
SZ_4 = Pt(14); SZ_XS4 = Pt(12); SZ_3 = Pt(16); SZ_2 = Pt(22); SZ_5 = Pt(10)

def mk_font(run, fn, sz, bold=False, cs=None):
    run.font.name = fn; run.font.size = sz; run.bold = bold
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None: rPr = OxmlElement('w:rPr'); run._r.insert(0, rPr)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), fn); rFonts.set(qn('w:ascii'), cs or fn); rFonts.set(qn('w:hAnsi'), cs or fn)
    rPr.insert(0, rFonts)

def add_p(text='', fn=SONGTI, sz=SZ_XS4, bold=False, align=None, indent=True, sb=0, sa=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.line_spacing = LINE_28
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    if indent: pf.first_line_indent = Pt(24)
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        r = p.add_run(text); mk_font(r, fn, sz, bold)
    return p

def blank(): add_p('', indent=False)

# ===== COVER PAGE =====
for _ in range(3): blank()
add_p('课 程 设 计', HEI, Pt(26), bold=True, align='center', indent=False, sb=20, sa=30)

for label in ['课程名称：', '设计题目：PWM直流调速系统设计', '',
              '学    院：', '专    业：', '年    级：',
              '学生姓名：', '学生学号：', '指导教师：', '日    期：']:
    if label == '':
        blank()
    elif label.startswith('设计题目'):
        add_p(label, HEI, SZ_4, bold=True, indent=False, sb=8, sa=8)
    else:
        add_p(label, indent=False, sb=6, sa=6)

blank()
add_p('教 务 处 制', align='center', indent=False, sb=30)

# ===== PAGE BREAK → TOC =====
doc.add_page_break()
add_p('目  录', HEI, SZ_3, bold=True, align='center', indent=False, sa=20)

toc = [
    ('第一章 绪论', '1'),
    ('  1.1 为什么要开设这门课', '1'),
    ('  1.2 运动控制系统的历史与发展', '1'),
    ('  1.3 为什么选择PWM直流调速系统', '2'),
    ('第二章 直流调速系统的工程设计方法', '3'),
    ('  2.1 工程设计方法的基本指导思想', '3'),
    ('  2.2 常用的典型系统', '4'),
    ('  2.3 双闭环直流调速系统', '5'),
    ('第三章 计算机Simulink仿真与计算', '7'),
    ('  3.1 调节器的设计参数计算', '7'),
    ('    3.1.1 电流调节器设计', '7'),
    ('    3.1.2 转速调节器设计', '9'),
    ('  3.2 主电路模型的参数计算', '12'),
    ('  3.3 Simulink仿真', '13'),
    ('    3.3.1 仿真模型及其参数设置', '13'),
    ('    3.3.2 仿真输出波形及其分析', '16'),
    ('第四章 具体硬件电路图设计与器件选型', '18'),
    ('  4.1 直流PWM放大器设计', '18'),
    ('    4.1.1 脉冲频率发生器的设计', '18'),
    ('    4.1.2 脉宽调制器的设计', '20'),
    ('  4.2 主电路', '22'),
    ('  4.3 调节器', '23'),
    ('    4.3.1 电流调节器', '23'),
    ('    4.3.2 转速调节器', '24'),
    ('  4.4 反馈模块设计', '24'),
    ('参考文献', '25'),
]
for item, page in toc:
    dots = '.' * (52 - len(item))
    add_p(f'{item}{dots}{page}', indent=False)

# ===== PAGE BREAK → BODY =====
doc.add_page_break()

data = json.load(open(r'C:\Users\lb\stock-quant\course_design_data_fixed.json', 'r', encoding='utf-8'))

for item in data:
    tag = item['t']; text = item['v']
    if tag == 's':  # Chapter heading
        add_p(text, HEI, SZ_3, bold=True, indent=False, sb=20, sa=10)
    elif tag == 'h':  # Section heading
        add_p(text, HEI, SZ_4, bold=True, indent=False, sb=12, sa=6)
    elif tag == 'h3':  # Subsection heading
        add_p(text, SONGTI, SZ_XS4, bold=True, indent=False, sb=8, sa=4)
    elif tag == 'b':  # Body text
        add_p(text)

# ===== REFERENCES =====
add_p('参考文献', HEI, SZ_3, bold=True, indent=False, sb=20, sa=10)
refs = [
    '[1] 陈伯时. 电力拖动自动控制系统——运动控制系统（第5版）[M]. 北京：机械工业出版社，2020.',
    '[2] 阮毅，杨影，陈伯时. 电力拖动自动控制系统——运动控制系统（第5版）习题解答及学习指导[M]. 北京：机械工业出版社，2020.',
    '[3] 王兆安，刘进军. 电力电子技术（第5版）[M]. 北京：机械工业出版社，2019.',
    '[4] Mohan N, Undeland T M, Robbins W P. Power Electronics: Converters, Applications, and Design (3rd Ed)[M]. New York: John Wiley & Sons, 2003.',
    '[5] 洪乃刚. 电力电子、电机控制系统的建模和仿真[M]. 北京：机械工业出版社，2019.',
    '[6] Bose B K. Modern Power Electronics and AC Drives[M]. Upper Saddle River: Prentice Hall, 2001.',
    '[7] 李发海，王岩. 电机与拖动基础（第5版）[M]. 北京：清华大学出版社，2018.',
    '[8] 胡寿松. 自动控制原理（第7版）[M]. 北京：科学出版社，2019.',
    '[9] Leonhard W. Control of Electrical Drives (3rd Ed)[M]. Berlin: Springer, 2001.',
    '[10] 薛定宇. 控制系统计算机辅助设计——MATLAB语言与应用（第3版）[M]. 北京：清华大学出版社，2012.',
]
for ref in refs:
    add_p(ref, indent=False, sb=2, sa=2)

# Save
outpath = r'C:/Users/lb/Desktop/PWM直流调速系统课程设计.docx'
doc.save(outpath)

chars = sum(len(p.text.replace(' ','').replace('\n','')) for p in doc.paragraphs)
print(f'Saved: {len(doc.paragraphs)}p ~{chars} chars -> {outpath}')
