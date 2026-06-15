"""Create TOC on design paper template."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

doc = Document()

# Page setup
for s in doc.sections:
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.17); s.right_margin = Cm(3.17)

SONGTI = '宋体'; HEI = '黑体'; KAITI = '楷体'
LINE_28 = Pt(28)

def mk_font(run, fn, sz, bold=False):
    run.font.name = fn; run.font.size = sz; run.bold = bold
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None: rPr = OxmlElement('w:rPr'); run._r.insert(0, rPr)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), fn); rFonts.set(qn('w:ascii'), fn); rFonts.set(qn('w:hAnsi'), fn)
    rPr.insert(0, rFonts)

def add_p(text='', fn=SONGTI, sz=Pt(12), bold=False, align=None, indent=False, sb=0, sa=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.line_spacing = LINE_28
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        r = p.add_run(text); mk_font(r, fn, sz, bold)
    return p

# ===== TOC =====
add_p('目  录', HEI, Pt(16), bold=True, align='center', sa=20)

toc = [
    ('第一章 绪论', '1'),
    ('  1.1 为什么要开设这门课', '1'),
    ('  1.2 运动控制系统的历史与发展', '1'),
    ('  1.3 为什么选择PWM直流调速系统', '2'),
    ('第二章 直流调速系统的工程设计方法', '3'),
    ('  2.1 工程设计方法的基本指导思想', '3'),
    ('  2.2 常用的典型系统', '4'),
    ('  2.3 双闭环直流调速系统', '5'),
    ('第三章 计算机Simulink仿真与计算', '7'),
    ('  3.1 调节器的设计参数计算', '7'),
    ('    3.1.1 电流调节器设计', '7'),
    ('    3.1.2 转速调节器设计', '9'),
    ('  3.2 主电路模型的参数计算', '12'),
    ('  3.3 Simulink仿真', '13'),
    ('    3.3.1 仿真模型及其参数设置', '13'),
    ('    3.3.2 仿真输出波形及其分析', '16'),
    ('第四章 具体硬件电路图设计与器件选型', '18'),
    ('  4.1 直流PWM放大器设计', '18'),
    ('    4.1.1 脉冲频率发生器的设计', '18'),
    ('    4.1.2 脉宽调制器的设计', '20'),
    ('  4.2 主电路', '22'),
    ('  4.3 调节器', '23'),
    ('    4.3.1 电流调节器', '23'),
    ('    4.3.2 转速调节器', '24'),
    ('  4.4 反馈模块设计', '24'),
    ('参考文献', '25'),
]

for item, page in toc:
    dots = '.' * (52 - len(item))
    add_p(f'{item}{dots}{page}')

outpath = r'C:\Users\lb\Documents\xwechat_files\wxid_kntl8qm95eag22_6034\msg\file\2026-06\（4）-7设计（论文）专用纸_目录.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
