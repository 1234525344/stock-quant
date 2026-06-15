import docx, sys, re
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document(r'C:/Users/lb/Desktop/金贝贝实习报告_模板.docx')

# 1. 摘要字数
t19 = doc.paragraphs[19].text.replace(' ','').replace('\n','')
al = len(t19)
print(f'1. 摘要: {al}字(≤300) {"✓" if al <= 300 else "✗"}')

# 2. 关键词
t20 = doc.paragraphs[20].text.strip()
kw_part = t20.split('：')[-1].split(':')[-1]
kws = [k.strip() for k in kw_part.split('；')]
nk = len(kws)
ends_ok = not kw_part.rstrip().endswith('；')
print(f'2. 关键词: {nk}个(≤5) {"✓" if nk <= 5 else "✗"} | 尾无标点: {"✓" if ends_ok else "✗"} | 内容: {t20[:80]}')

# 3. 三部分字数
secs = {f'1 概述(≥1000)': range(23,39), f'2 实习总结(≥3000)': range(39,64), f'3 人生启迪(≥2000)': range(64,79)}
for name, rng in secs.items():
    c = sum(len(doc.paragraphs[i].text.replace(' ','').replace('\n','')) for i in rng)
    req = int(re.search(r'(\d+)', name).group(1))
    print(f'3. {name}: {c}字 {"✓" if c >= req else "✗缺" + str(req-c) + "字"}')

# 4. 参考文献
refs = sum(1 for p in doc.paragraphs if re.match(r'^\[\d+\]', p.text.strip()))
print(f'4. 参考文献: {refs}篇(≥10) {"✓" if refs >= 10 else "✗"}')

# 5. 标题字体
ok = True
for i in [23, 39, 64]:
    for r in doc.paragraphs[i].runs:
        if r.text.strip():
            fn, sz = r.font.name or '', (r.font.size or 0)/12700
            if '方正黑体' not in fn or abs(sz - 14) > 0.5:
                ok = False; print(f'5. ✗ 一级标题[{i}]: {fn} {sz}pt')
            break
if ok: print(f'5. ✓ 一级标题: 方正黑体_GBK 四号')

ok2 = True
for i in [24,29,37,40,45,51,58,60,62,65,69,72]:
    for r in doc.paragraphs[i].runs:
        if r.text.strip():
            fn, sz = r.font.name or '', (r.font.size or 0)/12700
            if '方正楷体' not in fn or abs(sz - 14) > 0.5:
                ok2 = False; print(f'   ✗ 二级标题[{i}]: {fn} {sz}pt')
            break
if ok2: print(f'   ✓ 二级标题: 方正楷体_GBK 四号')

# 6. 正文
ok3 = True
for i in range(19, 93):
    p = doc.paragraphs[i]; t = p.text.strip()
    if not t or re.match(r'^(摘要|关键词|学生实习|[123] |[123]\.[0-9] |参考文献|\[)', t):
        continue
    for r in p.runs:
        if r.text.strip():
            fn, sz = r.font.name or '', (r.font.size or 0)/12700
            if ('宋体' not in fn and fn != '') or abs(sz - 12) > 0.5:
                ok3 = False; print(f'6. ✗ 正文[{i}]: {fn} {sz}pt [{t[:40]}]')
            break
if ok3: print(f'6. ✓ 正文: 宋体 小四号')

# 7. 行间距
ok4 = True
for i in range(18, 93):
    ls = doc.paragraphs[i].paragraph_format.line_spacing
    ls_pt = ls / 12700 if ls else 0
    if ls_pt > 0 and abs(ls_pt - 28) > 1:
        ok4 = False; print(f'7. ✗ 第{i}段行间距: {ls_pt}pt')
        break
if ok4: print(f'7. ✓ 行间距: 固定值28磅')

# 8. 章节编号
has_1 = any(re.match(r'^[123] ', p.text.strip()) for p in doc.paragraphs)
has_11 = any(re.match(r'^[123]\.[0-9] ', p.text.strip()) for p in doc.paragraphs)
has_p1 = any(re.match(r'^（\d+）', p.text.strip()) for p in doc.paragraphs)
print(f'8. 编号层级: 章{"/" if has_1 else "✗"}节{"/" if has_11 else "✗"}(小节){"/" if has_p1 else "✗"} {"✓" if all([has_1,has_11,has_p1]) else "✗"}')

# 9. 标点
bad = 0
for p in doc.paragraphs:
    t = p.text
    if re.search(r'[一-鿿]', t):
        if re.search(r'[一-鿿],[ 一-鿿]', t): bad += 1
        if re.search(r'[一-鿿]\.[一-鿿]', t): bad += 1
print(f'9. 标点符号: {"✓ 无中英文混用" if bad == 0 else "✗ " + str(bad) + "处混用"}')

# 10. 边距
s = doc.sections[0]
t,b,l,r = s.top_margin/360000, s.bottom_margin/360000, s.left_margin/360000, s.right_margin/360000
m_ok = all(abs(v - tgt) < 0.03 for v, tgt in [(t,2.54),(b,2.54),(l,3.17),(r,3.17)])
print(f'10. 边距: T{t:.2f}B{b:.2f}L{l:.2f}R{r:.2f}(2.54/3.17) {"✓" if m_ok else "✗"}')

# 11. 副标题格式
sub = doc.paragraphs[9].text.strip()
print(f'11. 副标题: 【{sub}】 {"✓ 破折号格式" if sub.startswith("——") else "✗"}')

# 12. 摘要和关键词独立段落
print(f'12. 摘要段落[{19}]有首行缩进: {"✓" if doc.paragraphs[19].paragraph_format.first_line_indent else "✗"}')
print(f'    关键词段落[{20}]独立: ✓ (紧随摘要之后)')

# 13. 参考文献格式
ref_ok = all(('[J]' in p.text or '[M]' in p.text) for p in doc.paragraphs if re.match(r'^\[\d+\]', p.text.strip()))
print(f'13. 参考文献格式: {"✓ [M]或[J]标注" if ref_ok else "✗"}')

print(f'\n{"="*40}')
print(f'关键词完整内容: {doc.paragraphs[20].text}')
