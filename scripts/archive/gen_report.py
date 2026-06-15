import json, sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.17)
    s.right_margin = Cm(3.17)

SONGTI = '宋体'
HEI = '黑体'
KAITI = '楷体'
LINE_28 = Pt(28)
SIZE_3 = Pt(16)
SIZE_4 = Pt(14)
SIZE_XS4 = Pt(12)
SIZE_2 = Pt(22)

def mk_font(run, fn, sz, bold=False):
    run.font.name = fn
    run.font.size = sz
    run.bold = bold
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._r.insert(0, rPr)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), fn)
    rFonts.set(qn('w:ascii'), fn)
    rFonts.set(qn('w:hAnsi'), fn)
    rPr.insert(0, rFonts)

def add_p(text, fn=SONGTI, sz=SIZE_XS4, bold=False, align=None, indent=True, sb=0, sa=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE_28
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(sb)
    pf.space_after = Pt(sa)
    if indent:
        pf.first_line_indent = Pt(24)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    mk_font(r, fn, sz, bold)
    return p

def blank():
    add_p('', indent=False)

# ===== COVER PAGE =====
for _ in range(3):
    blank()
add_p('存档号：                           学号：', indent=False)
for _ in range(3):
    blank()
add_p('石家庄铁路职业技术学院', HEI, SIZE_2, bold=True, align='center', indent=False)
add_p('岗 位 实 习 报 告', HEI, SIZE_2, bold=True, align='center', indent=False, sa=40)
for _ in range(2):
    blank()

for f in ['系    别          轨道交通系', '学生姓名', '学生班级', '专业名称', '指导教师',
          '实习单位          数汇引领(深圳)科技有限公司']:
    add_p(f, indent=False, sb=10, sa=10)

blank(); blank()
add_p('二〇二六年六月', align='center', indent=False, sb=40)

# ===== TOC =====
doc.add_page_break()
add_p('目  录', HEI, SIZE_3, bold=True, align='center', indent=False, sa=20)
toc = [
    ('一、实习单位介绍', '1'), ('  1. 公司概况', '1'), ('  2. 主营业务与技术方向', '1'),
    ('二、实习岗位及主要工作内容', '2'), ('  1. 实习岗位', '2'), ('  2. 主要内容', '2'),
    ('三、实习主要过程', '3'), ('  1. 入职安全培训', '3'), ('  2. 岗位技能培训', '3'),
    ('  3. 实践操作阶段', '4'), ('四、实习的主要收获和体会', '5'),
    ('参考文献', '7'), ('致  谢', '7'), ('岗位实习记录', '8'), ('岗位实习考核鉴定表', '9'),
]
for item, page in toc:
    add_p(f'{item}{chr(46) * (50 - len(item))}{page}', indent=False)

# ===== BODY =====
doc.add_page_break()

# Load content
p1 = json.load(open(r'C:\Users\lb\stock-quant\report_p1.json', 'r', encoding='utf-8'))
p2 = json.load(open(r'C:\Users\lb\stock-quant\report_p2.json', 'r', encoding='utf-8'))
all_parts = p1 + p2

is_refs = False
for tag, text in all_parts:
    if tag == 's':
        if text == 'refs':
            add_p('参考文献', HEI, SIZE_3, bold=True, indent=False, sb=16, sa=10)
            is_refs = True
            continue
        elif text == '致谢':
            is_refs = False
            add_p('致   谢', HEI, SIZE_3, bold=True, align='center', indent=False, sb=30, sa=16)
            continue
        is_refs = False
        add_p(text, HEI, SIZE_3, bold=True, indent=False, sb=16, sa=10)
    elif tag == 'h':
        add_p(text, HEI, SIZE_XS4, bold=True, indent=False, sb=6, sa=4)
    elif tag == 'b':
        if is_refs:
            add_p(text, indent=False, sb=2, sa=2)
        else:
            add_p(text)

# ===== 岗位实习记录 =====
doc.add_page_break()
add_p('岗位实习记录', HEI, SIZE_3, bold=True, align='center', indent=False, sa=20)

t1 = doc.add_table(rows=13, cols=2)
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
t1.style = 'Table Grid'
for row in t1.rows:
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(11)

hdr = t1.rows[0]
hdr.cells[0].text = '时  间'
hdr.cells[1].text = '内    容'
for cell in hdr.cells:
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            mk_font(r, HEI, SIZE_XS4, bold=True)

for i in range(1, 13):
    row = t1.rows[i]
    row.cells[0].text = '2025年  月  日－  月  日'
    row.cells[1].text = ''
    row.height = Cm(3)

blank()
add_p('企业指导教师签名：', indent=False)
add_p('年      月      日', indent=False, sb=6)
blank()
add_p('岗位实习记录一周至少一次，并请企业兼职教师审阅签名。不够可另附页。', KAITI, Pt(10), indent=False)

# ===== 岗位实习考核鉴定表 =====
doc.add_page_break()
add_p('岗位实习考核鉴定表', HEI, SIZE_3, bold=True, align='center', indent=False, sa=20)

t2 = doc.add_table(rows=19, cols=7)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
t2.style = 'Table Grid'

cells = t2.rows[0].cells; cells[0].merge(cells[1]); cells[2].merge(cells[3]); cells[4].merge(cells[6])
cells[0].text = '学生姓名'; cells[2].text = '系      轨道交通系'; cells[4].text = '班级'

cells = t2.rows[1].cells; cells[0].merge(cells[1]); cells[2].merge(cells[6])
cells[0].text = '实习单位'; cells[2].text = '数汇引领(深圳)科技有限公司'

cells = t2.rows[2].cells; cells[0].merge(cells[1]); cells[2].merge(cells[6])
cells[0].text = '实习岗位'; cells[2].text = '实习电气技术'

cells = t2.rows[3].cells; cells[0].merge(cells[1]); cells[2].merge(cells[6])
cells[0].text = '实习时间'; cells[2].text = '2026年  月  日  至  2026年  月  日，共计    天'

cells = t2.rows[4].cells; cells[0].merge(cells[1]); cells[2].merge(cells[4]); cells[5].merge(cells[6])
cells[0].text = '企业指导老师'; cells[2].text = ''; cells[5].text = '校内指导老师'

cells = t2.rows[5].cells
for c in cells[1:]: cells[0].merge(c)
cells[0].text = '个人总结\n（包括实习内容、实习任务完成情况、组织纪律、工作态度、实际技能的掌握及收获体会等）'

for ri in [6, 7]:
    cells = t2.rows[ri].cells
    for c in cells[1:]: cells[0].merge(c)
    cells[0].text = ''

cells = t2.rows[8].cells
cells[0].text = '实习单位\n鉴定意见'; cells[1].text = '等级\n项目'
cells[2].text = '优秀'; cells[3].text = '良好'; cells[4].text = '中等'; cells[5].text = '及格'; cells[6].text = '不及格'

for ri, item in enumerate(['敬业爱岗', '吃苦耐劳', '团队合作', '岗位技能']):
    cells = t2.rows[9 + ri].cells; cells[0].merge(cells[1]); cells[0].text = item

cells = t2.rows[13].cells; cells[0].merge(cells[1]); cells[2].merge(cells[6])
cells[0].text = '综合评价'; cells[2].text = '□优秀    □良好    □中等    □及格    □不及格'

cells = t2.rows[14].cells; cells[0].merge(cells[6])
cells[0].text = '企业指导教师：                     签 名（单位盖章）\n年    月    日'

cells = t2.rows[15].cells
cells[0].text = '实习报告\n鉴定意见'; cells[1].text = '等级\n项目'
cells[2].text = '优秀'; cells[3].text = '良好'; cells[4].text = '中等'; cells[5].text = '及格'; cells[6].text = '不及格'

for ri, item in enumerate(['实习记录', '实习报告']):
    cells = t2.rows[16 + ri].cells; cells[0].merge(cells[1]); cells[0].text = item

cells = t2.rows[18].cells; cells[0].merge(cells[1]); cells[2].merge(cells[6])
cells[0].text = '综合评价'; cells[2].text = '□优秀    □良好    □中等    □及格    □不及格'

# Extra rows
row = t2.add_row(); row.cells[0].merge(row.cells[6])
row.cells[0].text = '校内指导教师：\n\n\n年    月    日'

row = t2.add_row(); row.cells[0].merge(row.cells[6])
row.cells[0].text = '岗位实习\n综合成绩'

row = t2.add_row(); row.cells[0].merge(row.cells[6])
row.cells[0].text = '系审核意见\n\n负责人（签章）：\n                                              年    月    日'

row = t2.add_row(); row.cells[0].merge(row.cells[6])
row.cells[0].text = '注：表格打印出来后由相应人员填写，学生实习成绩按优、良、中、合格和不合格五级评定。'

# Font cleanup
for table in [t1, t2]:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if r.font.size is None:
                        mk_font(r, SONGTI, Pt(10), False)

# Save
outpath = r'C:/Users/lb/Desktop/岗位实习报告_数汇引领.docx'
doc.save(outpath)

total = sum(len(p.text.replace(' ','').replace('\n','')) for p in doc.paragraphs)
print(f'Saved: {len(doc.paragraphs)} paragraphs, ~{total} chars')
print(f'Path: {outpath}')
