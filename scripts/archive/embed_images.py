#!/usr/bin/env python
"""Embed diagrams into the course design report docx."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

path = r'C:/Users/lb/Desktop/电力系统工程课程设计_修正版.docx'
doc = Document(path)

IMGDIR = r'C:/Users/lb/stock-quant/diagrams'

def insert_img(para, img_name, width_inches=5.8):
    """Insert image after a given paragraph."""
    import os
    img_path = os.path.join(IMGDIR, img_name)
    if not os.path.exists(img_path):
        print(f'  WARN: {img_name} not found!')
        return

    ref = para._element
    parent = ref.getparent()
    idx = list(parent).index(ref)

    # Create image paragraph
    new_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    sp = OxmlElement('w:spacing'); sp.set(qn('w:line'), '360'); sp.set(qn('w:lineRule'), 'auto')
    pPr.append(sp)
    new_p.append(pPr)

    r = OxmlElement('w:r')
    drawing = OxmlElement('w:drawing')
    # Use simple inline shape approach
    # Add using python-docx API via a temporary paragraph
    parent.insert(idx + 1, new_p)

    # Now use python-docx to add the actual image to this paragraph
    # Get the document's paragraph object
    from docx.text.paragraph import Paragraph
    img_para = Paragraph(new_p, doc)

    run = img_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))

    return img_para

print('Embedding images...')

# Find paragraphs and insert images
# Use the actual para indices from the document
img_map = {}
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if '图3-1 照明分布图' in t or t.startswith('图3-1'):
        img_map['fig3_1_lighting.png'] = i
    elif '图4-1 供配电系统图' in t and '12回路' not in t and '12条' not in t:
        img_map['fig4_1_system.png'] = i
    elif '图4-2 空调等大功率' in t:
        img_map['fig4_2_hp_circuits.png'] = i
    elif '图4-3 卫生间浴霸' in t:
        img_map['fig4_3_bath.png'] = i
    elif '图4-4 主卧小功率' in t:
        img_map['fig4_4_bedroom.png'] = i
    elif t.endswith('6-1 灯具安装位置平面图') or '图6-1 灯具' in t:
        if 'fig6_1_2_lights.png' not in img_map:
            img_map['fig6_1_2_lights.png'] = i
    elif '图6-3 电气设备' in t:
        img_map['fig6_3_equipment.png'] = i
    elif '图6-4 插座安装平面图' in t:
        img_map['fig6_4_sockets.png'] = i

# Insert images (reverse order to maintain positions)
for img_name in ['fig6_4_sockets.png', 'fig6_3_equipment.png', 'fig6_1_2_lights.png',
                  'fig4_4_bedroom.png', 'fig4_3_bath.png', 'fig4_2_hp_circuits.png',
                  'fig4_1_system.png', 'fig3_1_lighting.png']:
    if img_name in img_map:
        p = doc.paragraphs[img_map[img_name]]
        insert_img(p, img_name, 5.5)
        print(f'  Inserted {img_name} after para[{img_map[img_name]}]')

outpath = r'C:/Users/lb/Desktop/电力系统工程课程设计_带图.docx'
import os
try: os.remove(outpath)
except: pass
doc.save(outpath)
print(f'\nSaved: {outpath}')
