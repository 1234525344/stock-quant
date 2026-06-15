# -*- coding: utf-8 -*-
"""DIRECTLY modify the original document's formatting, preserving ALL content + images."""
import sys, re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

# ===== LOAD ORIGINAL =====
SRC = r'C:/Users/lb/Documents/xwechat_files/wxid_kntl8qm95eag22_6034/msg/file/2026-06/岗位实习总结(1).docx'
DST = r'C:/Users/lb/Desktop/岗位实习总结_已排版.docx'

# Use unique name to avoid locks
import time
DST_BASE = r'C:/Users/lb/Desktop/岗位实习总结_已排版'
DST = DST_BASE + '.docx'
# Find unused name
for suffix in ['', '_1', '_2', '_3', '_4', '_5']:
    candidate = DST_BASE + suffix + '.docx'
    import os
    if not os.path.exists(candidate):
        DST = candidate
        break
shutil.copy2(SRC, DST)
print(f'Working on: {DST}')

doc = Document(DST)

FANG = '仿宋_GB2312'
HEI = '黑体'
LINE_PT = Pt(26)

def set_font(run, fn, sz, bold=False):
    """Set font on a single run."""
    run.font.name = fn
    run.font.size = sz
    run.bold = bold
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._r.insert(0, rPr)
    # Remove existing rFonts
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:eastAsia'), fn)
    rF.set(qn('w:ascii'), fn)
    rF.set(qn('w:hAnsi'), fn)
    rPr.insert(0, rF)

def set_para(p, fn=FANG, sz=Pt(14), bold=False, indent_pt=28, align=None, ls_pt=Pt(26)):
    """Set formatting on all runs in a paragraph."""
    pf = p.paragraph_format
    pf.line_spacing = ls_pt
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if indent_pt:
        pf.first_line_indent = Pt(indent_pt)
    else:
        pf.first_line_indent = None
    if align == 'C':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'L':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for r in p.runs:
        if r.text or True:  # process all runs, even empty
            set_font(r, fn, sz, bold)

def has_image(p):
    """Check if paragraph contains an image (drawing)."""
    for r in p.runs:
        drawings = r._r.findall(qn('w:drawing'))
        if drawings:
            return True
    return False

def is_heading(p):
    """Determine if paragraph is a heading."""
    t = p.text.strip()
    if not t: return False, False, False
    # Chapter-level
    if t in ['综述', '主体', '总结', '目录', '实习目标', '实习要求', '实习单位简介',
             '实习单位介绍', '企业文化', '主要产品', '实习岗位介绍', '岗位概念', '岗位职责',
             '实习内容', '入职报道', '岗前培训', '实习过程', '就业前景分析', '实习心得', '实习总结',
             '基本信息', '培训时间', '培训目的', '培训内容', '培训要求']:
        return True, True, False  # is_chapter, is_centered, is_sub
    # Section-level
    if re.match(r'^[一二三四五六七]、', t):
        return True, False, True  # heading, not centered, sub
    if re.match(r'^（[一二三四五六七]）', t):
        return True, False, True
    return False, False, False

def classify_para(p):
    """Classify paragraph: 'cover', 'toc', 'chapter', 'sub', 'body', 'img'"""
    t = p.text.strip()

    if has_image(p):
        return 'img'

    # Cover page = section 0
    # TOC = section 1, '目录' title, items with dots
    if re.match(r'^目\s*录$', t):
        return 'toc_title'
    if '.'*5 in t or '..' in t:
        return 'toc_item'

    if t in ['JIANGXI POLYTECHNIC UNIVERSITY', '机械工程']:
        return 'cover_title'
    if t == '机械工程学院岗位实习报告' or t.startswith('机械工程'):
        return 'cover_title'
    if t == '二〇二六年六月':
        return 'cover_date'

    # Chapter headings - centered, large
    if t in ['综述', '主体', '总结', '实习目标', '实习要求', '实习单位简介',
             '实习单位介绍', '企业文化', '主要产品', '实习岗位介绍', '岗位概念',
             '岗位职责', '实习内容', '入职报道', '岗前培训', '实习过程',
             '就业前景分析', '实习心得', '实习总结']:
        return 'chapter'

    # Section headings - bold
    if re.match(r'^[一二三四五六七]、', t):
        return 'sub_heading'
    if re.match(r'^（[一二三四五六七]）', t):
        return 'sub_heading'
    if re.match(r'^一\)', t):
        return 'sub_heading'
    if t in ['基本信息', '培训目的', '培训内容', '培训要求']:
        return 'sub_heading'

    # Sub (numbered)
    if re.match(r'^[123]\.\s', t):
        return 'sub_sub_heading'

    # Everything else is body text
    if len(t) > 5:
        return 'body'

    return 'other'

# ===== STEP 1: FIX SECTION MARGINS =====
# Target: Page=A4(210x297mm), Margins: T=35mm B=32mm L=28mm R=26mm
for s in doc.sections:
    s.page_width = Mm(210)
    s.page_height = Mm(297)
    s.top_margin = Mm(35)
    s.bottom_margin = Mm(32)
    s.left_margin = Mm(28)
    s.right_margin = Mm(26)
print('Margins fixed for all sections.')

# ===== STEP 2: ADD SECTION BREAK after cover, ADD PAGE NUMBERS =====
# The cover is section 0 (paras 0-4). We need to insert a section break at para 4
# to separate cover from TOC+body, then remove the old section breaks at para 22.

# Actually, let me first fix the root issue: delete the existing section breaks
# and create proper ones.

# Current structure:
# Section 0: paras 0-4 (cover) -> section break in para 4
# Section 1: paras 5-22 (TOC) -> section break in para 22
# Section 2: paras 23-305 (body) -> section break at end

# We need:
# Section 0: cover (no page number)
# Section 1: TOC + body (page numbers from 1)

# Let's merge section 1 and 2 into one section
# Remove section break in para 22 first
para22 = doc.paragraphs[22]
pPr22 = para22._element.find(qn('w:pPr'))
if pPr22 is not None:
    sectPr22 = pPr22.find(qn('w:sectPr'))
    if sectPr22 is not None:
        pPr22.remove(sectPr22)
        print('Removed section break at para 22.')

# Also need to merge section properties. Body section (now section 2, was section 1+2)
# should inherit from the last section.

# For section 0 (cover): no page number
s0 = doc.sections[0]
# Ensure no page number
for child in list(s0._sectPr):
    if child.tag == qn('w:pgNumType'):
        s0._sectPr.remove(child)
print('Section 0: cover, no page numbers.')

# For sections 1+ (TOC + body): page numbers from 1
# There should be 2 or 3 sections after removing the break
# Add page numbers to the last section (body)
body_section = doc.sections[-1]
# Remove any existing pgNumType
for child in list(body_section._sectPr):
    if child.tag == qn('w:pgNumType'):
        body_section._sectPr.remove(child)
# Add new
pg = OxmlElement('w:pgNumType')
pg.set(qn('w:start'), '1')
body_section._sectPr.append(pg)

# Add footer with PAGE field
body_section.different_first_page_header_footer = False
footer = body_section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
# Clear existing
for r in list(fp.runs):
    fp._p.remove(r._r)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run()
fr.font.name = FANG
fr.font.size = Pt(14)
# East-asian font for footer
rPr = fr._r.find(qn('w:rPr'))
if rPr is None: rPr = OxmlElement('w:rPr'); fr._r.insert(0, rPr)
rF = OxmlElement('w:rFonts')
rF.set(qn('w:eastAsia'), FANG); rF.set(qn('w:ascii'), FANG); rF.set(qn('w:hAnsi'), FANG)
rPr.insert(0, rF)
# PAGE field
for tag, val in [('begin','begin'), ('text',' PAGE '), ('end','end')]:
    if tag == 'text':
        el = OxmlElement('w:instrText')
        el.set(qn('xml:space'), 'preserve')
        el.text = val
    else:
        el = OxmlElement('w:fldChar')
        el.set(qn('w:fldCharType'), val)
    fr._r.append(el)
print('Page numbers added to body section.')

# ===== STEP 3: FORMAT EVERY PARAGRAPH =====
stats = {'cover': 0, 'toc': 0, 'chapter': 0, 'sub_heading': 0, 'sub_sub_heading': 0, 'body': 0, 'img': 0, 'other': 0}

for i, p in enumerate(doc.paragraphs):
    cls = classify_para(p)
    stats[cls] = stats.get(cls, 0) + 1

    if cls == 'cover_title':
        set_para(p, HEI, Pt(24), bold=True, align='C', indent_pt=0)
    elif cls == 'cover_date':
        set_para(p, FANG, Pt(14), bold=False, align='C', indent_pt=0)
    elif cls == 'toc_title':
        set_para(p, HEI, Pt(22), bold=True, align='C', indent_pt=0)
    elif cls == 'toc_item':
        # TOC items: 4号仿宋, no indent, left-aligned
        # Use tab + dot leader
        set_para(p, FANG, Pt(14), bold=False, indent_pt=0)
        # Add tab stop
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None: pPr = OxmlElement('w:pPr'); p._element.insert(0, pPr)
        old_tabs = pPr.find(qn('w:tabs'))
        if old_tabs is not None: pPr.remove(old_tabs)
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right'); tab.set(qn('w:pos'), '8844'); tab.set(qn('w:leader'), 'dot')
        tabs.append(tab); pPr.append(tabs)
        # Replace dots with tab
        for r in p.runs:
            if r.text and '..' in r.text:
                # Split name and page
                parts = r.text.split('..')
                if len(parts) >= 2:
                    name = parts[0].rstrip('.')
                    page = parts[-1].split()[0] if parts[-1].split() else ''
                    r.text = f'{name}\t{page}'
    elif cls == 'chapter':
        set_para(p, HEI, Pt(22), bold=True, align='C', indent_pt=0)
    elif cls == 'sub_heading':
        set_para(p, HEI, Pt(14), bold=True, indent_pt=0)
    elif cls == 'sub_sub_heading':
        set_para(p, HEI, Pt(14), bold=True, indent_pt=0)
    elif cls == 'body':
        set_para(p, FANG, Pt(14), bold=False, indent_pt=28)
    elif cls == 'img':
        # Image paragraphs: center, no indent
        set_para(p, FANG, Pt(14), indent_pt=0, align='C')
    elif cls == 'other':
        # Short labels, empty paragraphs: keep simple
        t = p.text.strip()
        if not t:
            set_para(p, FANG, Pt(14), indent_pt=0)
        else:
            set_para(p, HEI, Pt(14), bold=True, indent_pt=0)

print(f'\nFormatted: {stats}')
print(f'Total: {sum(stats.values())} paragraphs')

# ===== SAVE =====
import tempfile
tmp = DST + '.tmp'
doc.save(tmp)
# Replace original
os.remove(DST)
os.rename(tmp, DST)
print(f'Saved: {DST}')
