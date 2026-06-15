# -*- coding: utf-8 -*-
"""Format the original document in-place: margins, fonts, line spacing, page numbers."""
import sys, re, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Mm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

# Load FROM original path (don't copy - avoids lock)
SRC = r'C:/Users/lb/Documents/xwechat_files/wxid_kntl8qm95eag22_6034/msg/file/2026-06/岗位实习总结(1).docx'
doc = Document(SRC)

# Save to temp file, then move
import tempfile, shutil
fd, TMP = tempfile.mkstemp(suffix='.docx')
os.close(fd)

FANG = '仿宋_GB2312'
HEI = '黑体'

def set_run(run, fn, sz, bold=False):
    run.font.name = fn
    run.font.size = sz
    run.bold = bold
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None: rPr = OxmlElement('w:rPr'); run._r.insert(0, rPr)
    for old in list(rPr.findall(qn('w:rFonts'))):
        rPr.remove(old)
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:eastAsia'), fn); rF.set(qn('w:ascii'), fn); rF.set(qn('w:hAnsi'), fn)
    rPr.insert(0, rF)

def fmt_para(p, fn, sz, bold=False, indent_pt=28, align=None):
    pf = p.paragraph_format
    pf.line_spacing = Pt(26); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    if indent_pt: pf.first_line_indent = Pt(indent_pt)
    else: pf.first_line_indent = None
    if align: p.alignment = align
    for r in p.runs:
        set_run(r, fn, sz, bold)

# ===== FIX SECTIONS =====
for s in doc.sections:
    s.page_width = Mm(210); s.page_height = Mm(297)
    s.top_margin = Mm(35); s.bottom_margin = Mm(32)
    s.left_margin = Mm(28); s.right_margin = Mm(26)
print('1. Margins fixed')

# Remove section break at para 22 (merge TOC+body sections)
p22 = doc.paragraphs[22]
pPr22 = p22._element.find(qn('w:pPr'))
if pPr22 is not None:
    sp22 = pPr22.find(qn('w:sectPr'))
    if sp22 is not None: pPr22.remove(sp22)
print('2. Merged TOC+body sections')

# Page numbers on body section
body_s = doc.sections[-1]
for c in list(body_s._sectPr):
    if c.tag == qn('w:pgNumType'): body_s._sectPr.remove(c)
pg = OxmlElement('w:pgNumType'); pg.set(qn('w:start'), '1'); body_s._sectPr.append(pg)

# Cover section: no page numbers
cov_s = doc.sections[0]
for c in list(cov_s._sectPr):
    if c.tag == qn('w:pgNumType'): cov_s._sectPr.remove(c)

# Footer on body section
ft = body_s.footer; ft.is_linked_to_previous = False
fp = ft.paragraphs[0]
for r in list(fp.runs): fp._p.remove(r._r)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run()
fr.font.name = FANG; fr.font.size = Pt(14)
rPr = fr._r.find(qn('w:rPr'))
if rPr is None: rPr = OxmlElement('w:rPr'); fr._r.insert(0, rPr)
rF = OxmlElement('w:rFonts')
rF.set(qn('w:eastAsia'), FANG); rF.set(qn('w:ascii'), FANG); rF.set(qn('w:hAnsi'), FANG)
rPr.insert(0, rF)
for tp, vl in [('b','begin'),('t',' PAGE '),('e','end')]:
    if tp == 't':
        el = OxmlElement('w:instrText'); el.set(qn('xml:space'),'preserve'); el.text = vl
    else:
        el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'), vl)
    fr._r.append(el)
print('3. Page numbers set')

# ===== FORMAT ALL PARAGRAPHS =====
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()

    # Detect images
    has_img = any(r._r.findall(qn('w:drawing')) for r in p.runs)

    # Cover: para 0-4
    if i <= 4:
        if 'JIANGXI' in t or '机械' in t:
            fmt_para(p, HEI, Pt(22), True, 0)
        elif '二〇二六' in t:
            fmt_para(p, FANG, Pt(14), False, 0)
        else:
            fmt_para(p, FANG, Pt(14), False, 0)
        continue

    # TOC title (para 5)
    if i == 5:
        fmt_para(p, HEI, Pt(22), True, 0, WD_ALIGN_PARAGRAPH.CENTER)
        continue

    # TOC items (paras 6-22): add tab stops
    if 6 <= i <= 21:
        fmt_para(p, FANG, Pt(14), False, 0)
        # Replace dots with tabs and add tab stops
        for r in p.runs:
            if r.text and '......' in r.text:
                parts = r.text.split('.')
                name = parts[0].rstrip('.')
                page = ''
                for pt in reversed(parts):
                    if pt.strip().isdigit():
                        page = pt.strip()
                        break
                if name and page:
                    r.text = f'{name}\t{page}'
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None: pPr = OxmlElement('w:pPr'); p._element.insert(0, pPr)
        old_t = pPr.find(qn('w:tabs'))
        if old_t is not None: pPr.remove(old_t)
        tbs = OxmlElement('w:tabs')
        tb = OxmlElement('w:tab')
        tb.set(qn('w:val'), 'right'); tb.set(qn('w:pos'), '8844'); tb.set(qn('w:leader'), 'dot')
        tbs.append(tb); pPr.append(tbs)
        continue

    if i == 22:
        # Section break paragraph
        fmt_para(p, FANG, Pt(14), False, 0)
        continue

    # Body (paras 23+)
    if has_img:
        fmt_para(p, FANG, Pt(14), False, 0, WD_ALIGN_PARAGRAPH.CENTER)
    elif t in ['综述','主体','总结','实习目标','实习要求','实习单位简介',
               '实习单位介绍','企业文化','主要产品','实习岗位介绍','岗位概念',
               '岗位职责','实习内容','入职报道','岗前培训','实习过程',
               '就业前景分析','实习心得','实习总结']:
        fmt_para(p, HEI, Pt(22), True, 0, WD_ALIGN_PARAGRAPH.CENTER)
    elif re.match(r'^[一二三四五六七]、', t) or re.match(r'^（[一二三四五六七]）', t) or re.match(r'^一\)', t) or t in ['培训时间','培训目的','培训内容','培训要求','基本信息']:
        fmt_para(p, HEI, Pt(14), True, 0)
    elif re.match(r'^[123]\.\s', t):
        fmt_para(p, HEI, Pt(14), True, 0)
    elif len(t) > 3:
        fmt_para(p, FANG, Pt(14), False, 28)
    else:
        fmt_para(p, FANG, Pt(14), False, 0)

print(f'4. Formatted {len(doc.paragraphs)} paragraphs')

# ===== SAVE =====
doc.save(TMP)

# Move to desktop
DST = r'C:/Users/lb/Desktop/岗位实习总结_已排版.docx'
if os.path.exists(DST):
    try: os.remove(DST)
    except: DST = r'C:/Users/lb/Desktop/岗位实习总结_已排版_v1.docx'
shutil.move(TMP, DST)
print(f'\nSaved: {DST}')

# Verify
d2 = Document(DST)
s = d2.sections[0]
print(f'Verify margins: T={s.top_margin/36000:.0f} B={s.bottom_margin/36000:.0f} L={s.left_margin/36000:.0f} R={s.right_margin/36000:.0f}')
print(f'Verify sections: {len(d2.sections)}')
print(f'Verify paras: {len(d2.paragraphs)}')
# Check a body section font
for i in [27, 56, 117, 143, 284, 296]:
    p = d2.paragraphs[i]
    for r in p.runs:
        if r.text.strip():
            print(f'  [{i}] font={r.font.name} sz={(r.font.size or 0)/12700:.0f}pt [{r.text.strip()[:60]}]')
            break
