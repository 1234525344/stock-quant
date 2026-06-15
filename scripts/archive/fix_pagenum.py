#!/usr/bin/env python
"""Fix page numbers in the saved document by rebuilding sections properly."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Mm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

SRC = r'C:/Users/lb/Documents/xwechat_files/wxid_kntl8qm95eag22_6034/msg/file/2026-06/岗位实习总结_格式版.docx'
DST = SRC  # overwrite

# ===== READ EXISTING CONTENT =====
import docx as dx
src = dx.Document(SRC)

# Extract all paragraphs with their formatting
para_data = []
for i, p in enumerate(src.paragraphs):
    t = p.text
    fn = sz = bold = al = ind = None
    for r in p.runs:
        if r.text.strip():
            fn = r.font.name
            sz = r.font.size
            bold = r.bold
            break
    pf = p.paragraph_format
    al = p.alignment
    ind = pf.first_line_indent
    para_data.append({
        'text': t,
        'font': fn,
        'size': sz,
        'bold': bold,
        'align': al,
        'indent': ind
    })

print(f'Read {len(para_data)} paragraphs from source')

# ===== BUILD NEW DOCUMENT =====
doc = Document()
FANGSONG = '仿宋_GB2312'
HEI = '黑体'
LINE_26 = Pt(26)

def mk_run(run, fn, sz, bold=False):
    run.font.name = fn; run.font.size = sz; run.bold = bold
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None: rPr = OxmlElement('w:rPr'); run._r.insert(0, rPr)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), fn); rFonts.set(qn('w:ascii'), fn); rFonts.set(qn('w:hAnsi'), fn)
    rPr.insert(0, rFonts)

# ===== COVER SECTION =====
s0 = doc.sections[0]
s0.page_width = Mm(210); s0.page_height = Mm(297)
s0.top_margin = Mm(35); s0.bottom_margin = Mm(32)
s0.left_margin = Mm(28); s0.right_margin = Mm(26)

# Copy cover paragraphs (up to the TOC area - find '目  录')
cover_end = 0
for i, pd in enumerate(para_data):
    if pd['text'].strip() == '目  录':
        cover_end = i
        break

for pd in para_data[:cover_end]:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE_26; pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    if pd['indent']: pf.first_line_indent = pd['indent']
    if pd['align'] is not None: p.alignment = pd['align']
    if pd['text']:
        r = p.add_run(pd['text'])
        fn = pd['font'] or FANGSONG
        sz = pd['size'] or Pt(14)
        mk_run(r, fn, sz, pd['bold'] or False)

# ===== SECTION BREAK (cover → body) =====
doc.add_section()
s1 = doc.sections[1]
s1.page_width = Mm(210); s1.page_height = Mm(297)
s1.top_margin = Mm(35); s1.bottom_margin = Mm(32)
s1.left_margin = Mm(28); s1.right_margin = Mm(26)

# Page number start = 1
pgNumType = OxmlElement('w:pgNumType')
pgNumType.set(qn('w:start'), '1')
s1._sectPr.append(pgNumType)

# Footer with page number (四号仿宋, centered, Arabic)
s1.different_first_page_header_footer = False
footer = s1.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = fp.add_run()
run.font.name = FANGSONG; run.font.size = Pt(14)
rPr1 = run._r.find(qn('w:rPr'))
if rPr1 is None: rPr1 = OxmlElement('w:rPr'); run._r.insert(0, rPr1)
rFonts1 = OxmlElement('w:rFonts')
rFonts1.set(qn('w:eastAsia'), FANGSONG); rFonts1.set(qn('w:ascii'), FANGSONG)
rFonts1.set(qn('w:hAnsi'), FANGSONG)
rPr1.insert(0, rFonts1)

f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin'); run._r.append(f1)
it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '
run._r.append(it)
f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end'); run._r.append(f2)

# ===== COPY REMAINING PARAGRAPHS =====
for pd in para_data[cover_end:]:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE_26; pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    if pd['indent']: pf.first_line_indent = pd['indent']
    if pd['align'] is not None: p.alignment = pd['align']
    if pd['text']:
        r = p.add_run(pd['text'])
        fn = pd['font'] or FANGSONG
        sz = pd['size'] or Pt(14)
        mk_run(r, fn, sz, pd['bold'] or False)

# ===== SAVE =====
doc.save(DST)

total = sum(len(p.text.replace(' ','').replace('\n','')) for p in doc.paragraphs)
print(f'Saved: {len(doc.paragraphs)}p, ~{total} chars')
print(f'Sections: {len(doc.sections)}')
for i, s in enumerate(doc.sections):
    pg = s._sectPr.find(qn('w:pgNumType'))
    st = pg.get(qn('w:start')) if pg is not None else 'N/A'
    has_footer = False
    for fp in s.footer.paragraphs:
        for r in fp.runs:
            for c in r._r:
                if c.tag == qn('w:instrText'):
                    has_footer = True
    print(f"  Section {i}: pgNumType={st}, footer={'PAGE' if has_footer else 'none'}")
