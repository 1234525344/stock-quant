import sys
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_LINE_SPACING

sys.stdout.reconfigure(encoding='utf-8')

SONGTI = '宋体'
FZHT = '方正黑体_GBK'
FZKT = '方正楷体_GBK'
LINE_28 = Pt(28)

# Read paragraphs from source
import docx as dx
src = dx.Document(r'C:/Users/lb/Documents/xwechat_files/wxid_kntl8qm95eag22_6034/msg/file/2026-06/金贝贝实习报告002.docx')

paras = []
counting = False
for p in src.paragraphs:
    t = p.text.strip()
    if t == '学生实习报告正文':
        counting = True
        continue
    if counting and t:
        fn = None; sz = None; b = None
        for r in p.runs:
            if r.text.strip():
                fn = r.font.name; sz = r.font.size; b = r.bold
                break
        paras.append({'text': p.text, 'font': fn, 'size': sz, 'bold': b})

# Create new document
from docx import Document
doc = Document()

# Set page margins
for s in doc.sections:
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.17)
    s.right_margin = Cm(3.17)

def add_para(text, font_name, size, bold=False):
    """Add a paragraph with proper formatting."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.line_spacing = LINE_28
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Check if it's a title (no indent) or body (first-line indent)
    is_title = text.startswith(('1 ', '2 ', '3 ', '1.', '2.', '3.'))
    if not is_title:
        pf.first_line_indent = Pt(24)  # ~2 chars

    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.bold = bold

    # Set east-asian font
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._r.insert(0, rPr)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.insert(0, rFonts)

    return para

# Add all paragraphs
for p in paras:
    fn = p['font'] or SONGTI
    sz_pt = p['size'] / 12700 if p['size'] else 12
    sz = Pt(sz_pt)
    add_para(p['text'], fn, sz, p['bold'] or False)

# Handle Chinese font display for 方正 fonts
# In some environments, 方正 fonts need explicit east-asian mapping

outpath = r'C:/Users/lb/Desktop/金贝贝实习报告.docx'
doc.save(outpath)
print(f'Saved {len(paras)} paragraphs to {outpath}')

# Verify
doc2 = dx.Document(outpath)
total_chars = 0
for p in doc2.paragraphs:
    for r in p.runs:
        if r.text.strip():
            total_chars += len(r.text.replace(' ','').replace('\n',''))
            break
print(f'Total chars: {total_chars}')
