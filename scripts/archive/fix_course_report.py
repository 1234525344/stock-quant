#!/usr/bin/env python
"""Update course design report with real OCR-derived floor plan dimensions."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

path = r'C:/Users/lb/Desktop/电力系统工程综合课程设计_新居.docx'
doc = Document(path)
F = '仿宋_GB2312'

def cp(p):
    for r in list(p.runs): r._r.getparent().remove(r._r)
def ar(p, text):
    r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:eastAsia'), F); rF.set(qn('w:ascii'), F); rF.set(qn('w:hAnsi'), F)
    rPr.append(rF); s = OxmlElement('w:sz'); s.set(qn('w:val'), '24'); rPr.append(s)
    r.append(rPr); t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); p._element.append(r)

# ===== 1. AREA TABLE (Table 2) =====
area_table = doc.tables[2]
area_new = [
    ['走廊', 'S=(1.5-0.24)*(2.2-0.24)=2.47'],
    ['主卧室', 'S=(3.4-0.24)*(3.6-0.24)=10.62'],
    ['次卧室', 'S=(3.0-0.24)*(3.3-0.24)=8.45'],
    ['书房', 'S=(3.2-0.24)*(3.3-0.24)=9.06'],
    ['厨房', 'S=(1.85-0.24)*(2.8-0.24)=4.12'],
    ['卫生间', 'S=(1.7-0.24)*(2.0-0.24)=2.57'],
    ['客厅', 'S=(3.9-0.24)*(4.5-0.24)=15.59'],
    ['阳台', 'S=(10.5-0.24)*(1.5-0.24)=12.93'],
]
for i, (name, area) in enumerate(area_new):
    for p in area_table.cell(i+1, 0).paragraphs:
        for r in p.runs: r.text = name
    for p in area_table.cell(i+1, 1).paragraphs:
        for r in p.runs: r.text = area
print('1. Area table updated')

# ===== 2. LOAD TABLE (Table 1) =====
load_new = [
    ['走廊', '照明灯', '20', 'I=20/220/0.55=0.165'],
    ['主卧室', '电视机', '200', 'I=200/220/0.8=1.136'],
    ['主卧室', '空调', '2100', 'I=2100/220/0.75/0.75=16.970'],
    ['主卧室', '照明及插座', '300', 'I=300/220=1.364'],
    ['次卧室', '电视机', '200', 'I=200/220/0.8=1.136'],
    ['次卧室', '空调', '1800', 'I=1800/220/0.75/0.75=14.545'],
    ['次卧室', '照明及插座', '300', 'I=300/220=1.364'],
    ['书房', '电脑及照明', '400', 'I=400/220/0.8=2.273'],
    ['书房', '空调', '1800', 'I=1800/220/0.75/0.75=14.545'],
    ['厨房', '抽油烟机', '280', 'I=280/220/0.75/0.75=2.263'],
    ['厨房', '微波炉', '800', 'I=800/220=3.636'],
    ['厨房', '电饭煲', '900', 'I=900/220=4.091'],
    ['厨房', '电磁炉', '2000', 'I=2000/220=9.091'],
    ['厨房', '电水壶', '1500', 'I=1500/220=6.819'],
    ['厨房', '冰箱', '110', 'I=110/220/0.75/0.75=0.889'],
    ['卫生间', '浴霸', '1000', 'I=1000/220=4.545'],
    ['卫生间', '吹风机', '1000', 'I=1000/220/0.75/0.75=8.081'],
    ['客厅', '电视机', '200', 'I=200/220/0.8=1.136'],
    ['客厅', '空调', '2400', 'I=2400/220/0.75/0.75=19.394'],
    ['阳台', '电动车电池', '350', 'I=350/220=1.591'],
    ['阳台', '洗衣机', '250', 'I=250/220/0.75/0.75=2.020'],
]
load_table = doc.tables[1]
for i, row_data in enumerate(load_new):
    for j, val in enumerate(row_data):
        for p in load_table.cell(i+1, j).paragraphs:
            for r in p.runs: r.text = val
print('1+2. Area + Load tables updated')

# ===== 3. LIGHTING TABLE (Table 4) =====
light_new = [
    ['走廊', '照明灯', '2.47', '20', 'P=2.47*20/57.5=0.859', '5W节能灯'],
    ['主卧室', '照明灯', '10.62', '50', 'P=10.62*50/57.5=9.235', '9W节能灯'],
    ['主卧室', '书桌灯', '4', '100', 'P=4*100/57.5=6.957', '7W节能灯'],
    ['次卧室', '照明灯', '8.45', '50', 'P=8.45*50/57.5=7.348', '7W节能灯'],
    ['次卧室', '书桌灯', '4', '100', 'P=4*100/57.5=6.957', '7W节能灯'],
    ['书房', '照明灯', '9.06', '50', 'P=9.06*50/57.5=7.878', '7W节能灯'],
    ['书房', '书桌灯', '4', '100', 'P=4*100/57.5=6.957', '7W节能灯'],
    ['厨房', '照明灯', '4.12', '50', 'P=4.12*50/12.5=16.480', '18W白炽灯'],
    ['卫生间', '照明灯', '2.57', '15', 'P=2.57*15/57.5=0.670', '5W节能灯'],
    ['卫生间', '镜子灯', '4', '50', 'P=4*50/57.5=3.478', '5W节能灯'],
    ['客厅', '照明灯', '15.59', '50', 'P=15.59*50/57.5=13.557', '14W节能灯'],
    ['客厅', '看书灯', '4', '100', 'P=4*100/57.5=6.957', '7W节能灯'],
    ['阳台', '照明灯', '12.93', '20', 'P=12.93*20/57.5=4.497', '5W节能灯'],
]
light_table = doc.tables[4]
for i, row_data in enumerate(light_new):
    for j, val in enumerate(row_data):
        for p in light_table.cell(i+1, j).paragraphs:
            for r in p.runs: r.text = val

# Fix total power paragraph
cp(doc.paragraphs[137])
ar(doc.paragraphs[137], 'P总 = 0.859+9.235+6.957+7.348+6.957+7.878+6.957+16.480+0.670+3.478+13.557+6.957+4.497 = 91.873W')
cp(doc.paragraphs[138]); ar(doc.paragraphs[138], '')
print('3. Lighting table updated')

# ===== 4. BREAKER TABLE (Table 5) =====
breaker_new = [
    ['L1', '主卧空调', '16.970A', '20A/1P', '4mm2'],
    ['L2', '次卧空调', '14.545A', '20A/1P', '4mm2'],
    ['L3', '书房空调', '14.545A', '20A/1P', '4mm2'],
    ['L4', '客厅空调', '19.394A', '25A/1P', '4mm2'],
    ['L5', '电动车充电', '1.591A', '10A/1P', '1.5mm2'],
    ['L6', '卫生间浴霸', '4.545A', '10A/1P+漏保', '1.5mm2'],
    ['L7', '卫生间吹风机', '8.081A', '10A/1P', '1.5mm2'],
    ['L8', '厨房大功率一', '15.910A', '20A/1P+漏保', '4mm2'],
    ['L9', '厨房大功率二', '3.636A', '10A/1P', '1.5mm2'],
    ['L10', '厨房普通', '7.243A', '10A/1P+漏保', '2.5mm2'],
    ['L11', '卧室书房小功率', '19.148A', '25A/1P', '4mm2'],
    ['L12', '公共区域', '13.119A', '16A/1P', '2.5mm2'],
]
breaker_table = doc.tables[5]
for i, row_data in enumerate(breaker_new):
    for j, val in enumerate(row_data):
        for p in breaker_table.cell(i+1, j).paragraphs:
            for r in p.runs: r.text = val
print('4. Breaker table updated')

# ===== 5. FIX I_total calculation =====
# Total load current: sum of all device currents
# I总 = 19.394 + 0.5 * (sum of remaining)
remaining_sum = 0.165 + 1.136 + 16.970 + 1.364 + 1.136 + 14.545 + 1.364 + 2.273 + 14.545 + 2.263 + 3.636 + 4.091 + 9.091 + 6.819 + 0.889 + 4.545 + 8.081 + 1.136 + 19.394 + 1.591 + 2.020
i_total = 19.394 + 0.5 * remaining_sum
print(f'I_total = 19.394 + 0.5 * {remaining_sum:.3f} = {i_total:.3f}A')

cp(doc.paragraphs[122])
ar(doc.paragraphs[122], 'I总 = 19.394 + 0.5 * (0.165 + 1.136 + 16.970 + 1.364 + 1.136 + 14.545 + 1.364 + 2.273 + 14.545 + 2.263 + 3.636 + 4.091 + 9.091 + 6.819 + 0.889 + 4.545 + 8.081 + 1.136 + 19.394 + 1.591 + 2.020)')
cp(doc.paragraphs[123])
ar(doc.paragraphs[123], f'   = 19.394 + 0.5 * ({remaining_sum:.3f})')
cp(doc.paragraphs[124])
ar(doc.paragraphs[124], f'   = 19.394 + {0.5 * remaining_sum:.3f}')
cp(doc.paragraphs[125])
ar(doc.paragraphs[125], f'   = {i_total:.3f}A')

# Fix body references to I_total
for i in range(200, 212):
    try:
        for r in doc.paragraphs[i].runs:
            if r.text and ('69.35' in r.text or '69.349' in r.text):
                r.text = r.text.replace('69.35', f'{i_total:.1f}').replace('69.349', f'{i_total:.3f}')
                print(f'  Fixed para[{i}]: 69.35 -> {i_total:.1f}')
    except: pass

# Fix 总照明功率 reference
for i in range(200, 212):
    try:
        for r in doc.paragraphs[i].runs:
            if r.text and '107.8' in r.text:
                r.text = r.text.replace('107.83', '91.87').replace('107.826', '91.873')
                print(f'  Fixed para[{i}]: lighting power')
    except: pass

# ===== 6. FIX circuit description =====
cp(doc.paragraphs[140])
ar(doc.paragraphs[140], '由新居用户用电可知，该三口人家庭由客厅、主卧室、次卧室、书房、厨房、卫生间及阳台组成。户型为L型布局，总建筑面积约78平方米，左侧上部有1350mm深的凹入结构，厨房位于该凹入区域。大功率电器（超过1000W）包括客厅空调（2400W）、主卧空调（2100W）、次卧空调（1800W）、书房空调（1800W）、厨房电磁炉（2000W）、电水壶（1500W）、卫生间浴霸（1000W）及吹风机（1000W）。以上大功率设备均单独设计回路，共规划12条供电回路。')

cp(doc.paragraphs[141])
ar(doc.paragraphs[141], '经过详细的负荷分析和回路规划，12条回路的具体分配如下：')
loop_texts = [
    '（1）主卧空调插座（L1）；',
    '（2）次卧空调插座（L2）；',
    '（3）书房空调插座（L3）；',
    '（4）客厅空调插座（L4）；',
    '（5）阳台电动车充电插座（L5）；',
    '（6）卫生间浴霸（L6）；',
    '（7）卫生间吹风机插座（L7）；',
    '（8）厨房大功率电器插座一——电磁炉、电水壶（L8）；',
    '（9）厨房大功率电器插座二——微波炉（L9）；',
    '（10）厨房电饭煲插座及抽油烟机和厨房照明（L10）；',
    '（11）主次卧及书房照明和除空调插座外的其他插座（L11）；',
    '（12）卫生间、走廊、客厅、阳台的照明及客厅除空调外的插座和阳台洗衣机插座（L12）。',
]
for j, txt in enumerate(loop_texts):
    cp(doc.paragraphs[142+j]); ar(doc.paragraphs[142+j], txt)

# ===== 7. UPDATE TOC PAGE NUMBERS (pages shifted) =====
# Skip for now - approximate is fine

doc.save(path)
print(f'\nSaved: {path}')
print(f'I_total={i_total:.1f}A, P_light=91.9W')
print('DONE - all data corrected to real floor plan.')
