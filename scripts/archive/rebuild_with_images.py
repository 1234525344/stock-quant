#!/usr/bin/env python
"""Final rebuild WITH images preserved from original."""
import sys, re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Mm, Cm, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

TITLE_FONT = '黑体'
BODY_FONT = '仿宋_GB2312'
MEDIA_DIR = r'C:\Users\lb\stock-quant\media_extracted'

doc = Document()

def mk(run, fn, sz, bold=False):
    run.font.name = fn; run.font.size = sz; run.bold = bold
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None: rPr = OxmlElement('w:rPr'); run._r.insert(0, rPr)
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:eastAsia'), fn); rF.set(qn('w:ascii'), fn); rF.set(qn('w:hAnsi'), fn)
    rPr.insert(0, rF)

def P(text='', fn=BODY_FONT, sz=Pt(14), bold=False, align=None, indent_pt=0, sb=0, sa=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(26); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    if indent_pt: pf.first_line_indent = Pt(indent_pt)
    if align == 'C': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        r = p.add_run(text); mk(r, fn, sz, bold)
    return p

def img(img_name, width_inches=5.5):
    """Add an image paragraph (centered)."""
    import os
    img_path = os.path.join(MEDIA_DIR, img_name)
    if not os.path.exists(img_path):
        print(f'  WARNING: {img_name} not found!')
        return P(f'[图片: {img_name}]', indent_pt=0, align='C')
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(26); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(img_path, width=Inches(width_inches))
    return p

def title(text): return P(text, TITLE_FONT, Pt(22), True, 'C', sb=6, sa=10)
def subhead(text): return P(text, TITLE_FONT, Pt(14), True, sb=10, sa=4)
def body(text): return P(text, BODY_FONT, Pt(14), indent_pt=28)
def body0(text): return P(text, BODY_FONT, Pt(14))
def blank(): return P('')

# ===== COVER =====
s0 = doc.sections[0]
for a, v in [('page_width', Mm(210)), ('page_height', Mm(297)),
    ('top_margin', Mm(35)), ('bottom_margin', Mm(32)),
    ('left_margin', Mm(28)), ('right_margin', Mm(26))]:
    setattr(s0, a, v)

# Cover image (top of page - from original para 0)
img('image1.png', 2.0)
for _ in range(2): blank()
P('JIANGXI POLYTECHNIC UNIVERSITY', TITLE_FONT, Pt(15), True, 'C')
P('机械工程学院岗位实习报告', TITLE_FONT, Pt(24), True, 'C', sb=4, sa=40)
for _ in range(4): blank()
P('二〇二六年六月', BODY_FONT, Pt(14), align='C', sb=30)

# ===== SECTION BREAK =====
doc.add_section()
s1 = doc.sections[1]
for a, v in [('page_width', Mm(210)), ('page_height', Mm(297)),
    ('top_margin', Mm(35)), ('bottom_margin', Mm(32)),
    ('left_margin', Mm(28)), ('right_margin', Mm(26))]:
    setattr(s1, a, v)

# Page number
pg = OxmlElement('w:pgNumType'); pg.set(qn('w:start'), '1'); s1._sectPr.append(pg)
ft = s1.footer; ft.is_linked_to_previous = False
fp = ft.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(); fr.font.name = BODY_FONT; fr.font.size = Pt(14)
rp = fr._r.find(qn('w:rPr'))
if rp is None: rp = OxmlElement('w:rPr'); fr._r.insert(0, rp)
rf = OxmlElement('w:rFonts')
rf.set(qn('w:eastAsia'), BODY_FONT); rf.set(qn('w:ascii'), BODY_FONT); rf.set(qn('w:hAnsi'), BODY_FONT)
rp.insert(0, rf)
for t, v in [('begin','begin'), ('instr',' PAGE '), ('end','end')]:
    el = OxmlElement('w:fldChar') if t != 'instr' else OxmlElement('w:instrText')
    if t == 'instr': el.set(qn('xml:space'), 'preserve'); el.text = v
    else: el.set(qn('w:fldCharType'), v)
    fr._r.append(el)

# ===== TOC with tab stops =====
title('目  录')
toc = [
    ('综述', '1'), ('  实习单位基本情况', '1'), ('  实习岗位描述', '2'),
    ('主体', '3'), ('  职业素养与社会能力提升', '3'),
    ('  实习内容、收获与技能掌握', '4'),
    ('  发现、分析与解决实际问题的能力简述', '6'),
    ('  专业技能与岗位要求的匹配度与差距', '7'),
    ('  对专业课程与实习安排的建议', '8'),
    ('  今后努力方向', '9'), ('总结', '10'),
]
for item_name, page_num in toc:
    p = P(indent_pt=0)
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None: pPr = OxmlElement('w:pPr'); p._element.insert(0, pPr)
    old = pPr.find(qn('w:tabs'))
    if old is not None: pPr.remove(old)
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '8844')
    tab.set(qn('w:leader'), 'dot')
    tabs.append(tab); pPr.append(tabs)
    for r in p.runs: r.text = ''
    if p.runs: p.runs[0].text = f'{item_name}\t{page_num}'
    else: r = p.add_run(f'{item_name}\t{page_num}'); mk(r, BODY_FONT, Pt(14))

doc.add_page_break()

# ===== 综述 =====
title('综述')
subhead('实习单位基本情况')
body('东莞市浩镒晟五金电子制品有限公司，地址位于广东省东莞市官井头水库工业一路8号，是一家专业从事五金电子制品研发、生产与销售于一体的制造型企业。公司经营范围涵盖CNC机加工、锌铝压铸及压铸模具制造、冲压、脱蜡铸造及模具制造、钣金加工、粉末冶金制品、塑胶制品等多元化金属加工领域。')
body('公司秉承"严谨务实，敬业精进；团结拼搏，高效担当"的企业精神，建立了从产品研发设计、模具制造、精密加工、表面处理到品质检验的完整生产体系。公司拥有先进的CNC加工中心、压铸机、冲压设备、抛光及电镀生产线，配备了三次元测量仪、硬度计、盐雾试验机等完善的检测设备。公司注重技术创新和人才培养，逐步构建了标准化、精细化的生产管理模式，是珠三角地区五金电子加工行业具有代表性的制造企业之一。')

# === Product images ===
subhead('主要产品')
# image2-7: individual product photos (脱蜡铸造, 粉末冶金, CNC, 锌合金, 铝合金)
body('脱蜡铸造类不锈钢及碳素钢制品：')
img('image2.jpeg', 4.0)
body('粉末冶金类制品：')
img('image3.jpeg', 4.0)
body('CNC车床类精密零部件：')
img('image4.jpeg', 4.0)
body('锌合金压铸类制品：')
img('image5.jpeg', 4.0)
body('铝合金压铸类制品：')
img('image6.jpeg', 4.0)
body('产品广泛应用于消费电子、汽车配件、通信设备、医疗器械等行业。')
img('image7.jpeg', 5.0)

# Process photo
body('公司在生产过程中严格遵循行业质量标准，以下为生产车间及加工过程实拍：')
img('image8.jpeg', 5.0)

# Appearance samples (7 images: image9.png through image17.jpeg)
body('外观件样品展示：')
img('image9.png', 5.5)
img('image10.png', 5.5)
img('image11.png', 5.5)
img('image12.png', 5.5)
img('image13.jpeg', 5.5)
img('image14.jpeg', 5.5)
# image15-17 are smaller
img('image15.jpeg', 4.0)
img('image16.jpeg', 4.0)
img('image17.jpeg', 4.0)

subhead('实习岗位描述')
body('我的实习岗位为工程助理，隶属于公司工程部。工程助理是制造业工程技术团队中承上启下的基础性核心辅助岗位，是衔接工程师与生产一线的重要纽带。本次实习立足于五金电子制品加工制造行业，以工程技术辅助工作为核心，依托五金零件加工、金属材质处理、产品质检管控、工程资料整理等专业工作场景，将机械工程、五金加工的基础理论知识结合企业实际生产流程落地应用，熟悉五金电子产品生产工艺、工程作业标准以及工厂运营管理模式，掌握工程助理基础实操技能。')
body('具体岗位职责包括以下六个方面：一是工程资料管理职责——负责工程图纸、生产工艺文件、技术标准文档的整理、归档、录入工作，规范保管零件加工图纸、质检标准、工艺参数等资料，制作工程台账、BOM物料标签、数据报表，精准记录零件规格、材质参数、加工数据，保证工程资料完整清晰可追溯；二是生产现场辅助职责——跟进五金零件生产加工流程，熟悉铝材、不锈钢、锌合金等金属材料加工工艺，协助工程师巡查生产车间，观察冲压、加工、成型等生产工序，记录生产过程中的工艺问题和设备运行情况，及时反馈现场异常；三是产品质检管控职责——配合质检部门完成五金成品及半成品的抽检工作，对照行业技术标准和产品要求检测零件外观、尺寸及表面处理效果，记录质检数据，整理不良品信息，分类标注瑕疵问题，辅助优化产品加工品质；四是物料与成本辅助核算——根据零件尺寸及材质进行重量和用料核算，统计物料消耗情况，整理原材料及半成品物料清单，协助完成物料盘点，为生产成本核算和物料管控提供精准数据支撑；五是技术跟进与学习职责——学习五金金属表面处理工艺、零件加工工艺、行业技术规范等专业知识，协助工程师优化简易工艺流程，整理技术改良记录，配合完成样品试制和参数调试等辅助工作；六是部门沟通协同职责——对接生产部、质检部、工程部，传递工艺通知、质检结果、生产整改要求，做好日常工作对接记录，协助完成部门之间资料流转和信息同步，保障生产工程工作有序推进。')

# ===== 主体 =====
title('主体')
subhead('职业素养与社会能力提升')

subhead('1. 规章制度与劳动纪律')
body('实习期间，我严格遵守公司厂区及工程部各项规章制度，服从工作安排，认真完成各项辅助工作，做到工作按时、保质、保量完成，不敷衍、不拖沓。每日提前到岗参加部门早会，了解当日工作任务安排，无迟到、早退、旷工现象。工作中秉持严谨认真的态度，对待尺寸测量、品质判定、数据记录等工作反复核查，减少工作失误，培养工程师精准严谨的工作素养。遇到生产异常、产品不良、工艺问题时，第一时间做好记录并向主管反馈，不擅自改动生产工艺和图纸参数，严格遵循工程修改流程。合理规划每日工作，分清轻重缓急，高效完成文职整理、现场跟进、样品检测等多项任务，提升工作效率。同时严格遵守公司保密制度，不私自外传、拷贝公司产品图纸、工艺参数、客户产品资料等核心技术资料。')

subhead('2. 职业道德')
body('在工程助理岗位上，我深刻认识到工程技术工作容不得马虎——一个数据错误可能导致批量生产报废，一个型号标错可能造成采购损失。因此，我始终坚守诚实守信、认真负责的职业道德底线：所有工程数据和检测数据真实客观记录，绝不篡改、编造数据；图纸、文件、报表整理规范，分类存档，保障工程资料完整性。对于工作中发现的任何技术疑问，坚持向主管或老员工请教确认后再执行，绝不凭感觉擅自处理。在跨部门配合工作中，注重团队协作，配合工程师、车间工作人员完成生产改良、样品优化、品质整改等工作，虚心接受工作指导，保持良好职业心态。')

subhead('3. 社会适应能力与存在不足')
body('初入制造型企业，我从校园相对宽松的环境切换到工厂严格的生产办公节奏，经历了从陌生到适应的过程。通过岗前培训和日常工作的锻炼，我较快地融入了工程部团队，适应了制造业标准化、精细化的工作模式。在日常工作中，我注重与生产车间、质检部、仓储部等部门人员的沟通，清晰传递工程要求、生产整改意见，逐步提升了跨部门协作和工作对接能力。在协助项目跟进过程中，学会了主动沟通、主动跟进、主动反馈，改变了学生时代的被动等待习惯。')
body('不足之处在于，实习初期对厂区生产流程不够熟悉，在跨部门沟通时偶尔因不了解对方的具体工作内容而导致沟通效率偏低，后期通过主动深入车间学习和向老员工请教，情况逐步改善。在应对突发的生产异常或紧急项目任务时，临场应变能力还有待进一步加强。')

subhead('实习内容、收获与技能掌握')
subhead('1. 主要工作内容')
body('（1）工程图纸与模型优化工作。日常承接五金电子产品的三维模型修改工作，根据客户需求和生产加工难点优化模型结构，修正模型尺寸偏差和结构不合理等问题，适配车间加工工艺。完成修改后，按照行业标准出具二维工程图纸，精准标注产品尺寸、公差范围、材质信息及表面处理方式，确保图纸清晰规范，可直接用于车间生产加工。工作中严格核对图纸数据，反复校验模型结构，杜绝图纸错误导致的生产偏差。')
body('（2）物料台账编制与产品确认。负责编制和更新BOM表及物料清单，统计工程项目所需原材料、五金配件、加工辅料，详细标注物料规格、材质、数量、采购编号，做好物料出入库登记，实时更新物料台账，保障物料管控清晰。同时配合项目负责人确认项目产品，核对产品图纸、样品、加工要求，确认产品材质、外观、尺寸标准，审核样机加工方案，为批量生产做好前期准备。')
body('（3）工程项目跟进管理。全程跟进分配的五金电子零部件项目，对接生产车间跟进加工进度，记录项目生产节点，及时反馈生产过程中的图纸修改和工艺调整问题。协调工程部与生产部的工作衔接，跟进样机打样、批量加工、表面处理全流程。针对氧化、喷砂加工工序，确认表面处理效果，核对外观色泽和粗糙度是否符合客户标准，及时整改不合格样机。')
body('（4）跨部门辅助配合工作。在完成本职工程工作之余，积极协助其他部门开展工作。协助生产部完成项目成品装配，按照装配流程组装五金零部件，核对配件适配性，排查装配卡顿和配件不符等问题。协助品检部开展品检工作，参照工程图纸和质检标准，检测产品尺寸、外观、表面处理质量，记录不合格产品问题并上报。配合加工车间完成镭雕辅助工作，核对镭雕图案、文字和位置参数，检查镭雕清晰度，保障产品标识符合出厂要求。')

subhead('2. 收获与体会')
body('（1）工程绘图与办公技能从生疏到熟练。从一开始对三维建模、二维出图不熟练，到能够独立完成模型修改、工程图纸输出，我真正掌握了工程绘图的规范和行业标准。通过编制BOM表和物料清单，养成了严谨细致、反复核对的工作习惯，明白了工程资料容不得马虎。同时规范制作工程报表、质检记录表、物料统计表的能力也有了实质性的提升。')
body('（2）制造工艺认知从课本到实践的跨越。在学校课堂上学习机械制造工艺、金属材料学等课程时，更多停留在书本概念。通过在日常工作中接触铝材、锌合金、不锈钢等金属原材料的实际加工应用，学习压铸、冲压、抛光、电镀、氧化、喷砂等表面处理工艺，我对材料特性和工艺适配性的理解从书本走向了生产现场。通过参与成品装配、品检、镭雕等实操工作，我真正了解五金电子产品从设计、打样、表面处理、装配到质检量产的完整流程。')
body('（3）职场思维方式从学生到职业人的转变。实习让我学会了主动沟通、主动跟进、主动承担。项目跟进过程中需要对接工程师、生产、品检、业务等多个部门，锻炼了协调能力和执行力。协助其他部门工作的过程也让我懂得团队配合的重要性——任何岗位都是公司整体流程中不可或缺的一环。')

subhead('3. 技能掌握程度')
body('通过五至六个月的实习，我在以下方面取得了较为扎实的进步：工程绘图方面，能够独立完成中等复杂程度的3D模型修改和2D工程图纸出图，掌握尺寸标注、公差标注、表面处理标注等制图规范；物料管理方面，能够独立编制BOM表、物料清单和物料台账，掌握物料编码规范和ERP系统基本操作；工艺认知方面，熟悉了常见五金加工工艺流程和表面处理工艺，了解不同材质产品的加工特点和品质要求；品质检测方面，掌握卡尺、千分尺、高度规等常规检测工具的使用，能够参照质检标准完成基础产品检测；办公技能方面，能够规范制作各类工程报表和数据统计表格。但在精密公差分析、复杂模具结构理解、独立解决现场工艺难题等方面仍存在较大差距，需要持续学习和积累。')

subhead('发现、分析与解决实际问题的能力简述')
body('问题发现。在跟进某批次锌合金压铸外观件项目时，我发现经喷砂和阳极氧化处理后，部分产品的表面出现了局部色泽不均的问题，与客户确认的标准样件存在肉眼可见的色差。该批次产品数量约200件，不良率约为15%。若不能及时解决，将影响整批产品的按时交付。')
body('问题分析。我首先对不良品进行了分类统计，发现色差主要集中在产品的边缘和转角区域。向车间喷砂操作工了解后得知，该批次产品在喷砂工序中存在挂具装载密度不均匀的情况。经查阅该产品的工艺文件和历史生产记录，并与主管工程师讨论，初步判断原因有三个可能性：一是挂具装载过密导致部分区域的喷砂覆盖率不足，二是阳极氧化槽液浓度或温度有波动，三是产品在氧化前的水洗工序中未充分清洁。')
body('解决方案。在主管工程师的指导下，我协助采取了以下措施：第一，优化喷砂挂具装载布局，将每挂负载量从原来的24件减少至18件，确保每个产品表面都能均匀受砂；第二，复查并标定阳极氧化槽液的浓度和温度参数，确认工艺参数在标准范围内；第三，加强氧化前处理工序的水洗和除油检查。整改方案实施后，对后续两个批次共400件产品进行了全检，表面色泽均匀度明显改善，不良率从15%降至2%以下。')
body('反思总结。这次经历让我深刻认识到，表面处理质量问题的排查需要从多环节入手——不是简单地更换工艺参数就能解决，而是要系统地追溯每个可能的影响因素。同时，问题处理的效率取决于数据和事实的充分程度。此外与一线操作人员的沟通是获取第一手信息的关键。')

subhead('专业技能与岗位要求的匹配度与差距')
subhead('1. 适应情况')
body('通过在校期间学习的机械制图、机械制造基础、金属材料学、互换性与测量技术等专业课程，我具备了识读工程图纸、理解尺寸公差、区分金属材质的基础能力，这些知识为快速上手工程助理岗位工作提供了重要支撑。在校期间培养的学习习惯和自律意识也帮助我较快地适应了企业的工作节奏和管理制度。')

subhead('2. 存在的差距')
body('（1）复杂图纸的理解和绘制能力不足。对于涉及多零件装配关系、复杂剖面视图和精密公差标注的工程图纸，独立阅读和准确理解的能力还有欠缺。（2）工艺分析深度不够。对五金加工和表面处理工艺的了解目前主要在操作层面，对于工艺参数的优化调整、不同材料之间工艺兼容性的分析还需要大量实践积累。（3）独立解决现场问题的能力薄弱。虽然在实习期间参与了一些问题的分析和处理，但大多是在主管工程师的指导下完成的。（4）项目管理能力空白。对于项目的计划编制、进度管控、成本核算和风险预警等管理层面的工作还缺乏系统了解。（5）行业知识面偏窄。对于五金电子行业的技术标准体系、新材料发展趋势、智能制造技术的应用等方面了解不足。')

subhead('对专业课程与实习安排的建议')
subhead('1. 课程设置与安排')
body('（1）强化工程制图与三维建模实训环节，建议在机械制图课程中增加基于企业实际产品的实训项目。（2）增设制造工艺综合实训课程，帮助学生建立从设计到制造的全流程认知。（3）加强BOM和物料管理相关教学内容。（4）增加品质管理和检测技术的教学内容。')

subhead('2. 教师教学')
body('（1）注重实际案例教学，多采用制造业企业的实际产品和工艺案例进行教学，拉近课堂教学与企业实践的距离。（2）强调规范和标准意识，在教学过程中反复强调工程标准、工艺规范和质量标准的重要性。')

subhead('3. 顶岗实习安排')
body('（1）实习前期增设集中岗前培训，涵盖行业认知、岗位基础技能和安全规范等内容。（2）加强实习过程的阶段性评估，建议在实习中期安排一次集中交流和评估。（3）建立更为系统的实习培养梯度，建议企业安排从观摩到协助再到独立的渐进式培养路径。')

subhead('今后努力方向')
body('基于对自身不足的清醒认识和制造业工程技术岗位的发展要求，我明确了今后的努力方向。')
body('第一，持续提升专业绘图和软件操作能力。进一步学习SolidWorks、AutoCAD等工程软件的高级功能，达到独立完成中等复杂程度产品设计修改的水平。')
body('第二，深入学习制造工艺专业知识。系统学习五金加工、表面处理、模具设计等方面的专业技术知识，持续关注新材料、新工艺、新技术在制造业中的应用发展。')
body('第三，加强现场实践经验积累。在今后的工作中扎根生产一线，主动参与更多的产品开发和生产改善项目，在实战中积累经验、锻炼独立分析和解决工程问题的能力。')
body('第四，培养项目管理和团队协作能力。学习基本的项目管理知识和方法，提升在多任务并行环境下的工作规划和组织能力。')
body('第五，坚持终身学习的职业态度。工程技术和制造业在不断发展进步，我将始终保持学习的热情和进取心，朝着专业工程技术方向稳步发展。')

# ===== 总结 =====
title('总结')
body('为期五至六个月的东莞浩镒晟五金电子制品有限公司工程助理岗位实习，是我从机械工程专业学生向工程技术从业者转型的重要起点，更是一段收获满满、感悟深刻的职业启蒙经历。')
body('回顾这段实习经历，我从一个对五金电子制造行业仅有模糊认知的在校学生，成长为能够独立完成3D模型修改、2D工程图纸出图、BOM物料编制、工程项目跟进、表面处理工艺确认、产品品质检验等多项岗位核心工作的工程助理。在专业技能层面，三维建模和工程绘图从生疏到熟练，物料台账从不知从何下手到规范编制，工艺流程从课本概念到亲身实践。在职业素养层面，我养成了严谨细致、按标准做事、主动沟通跟进的工作习惯，初步完成了从学生心态到职业心态的转变。')
body('同时，这段实习经历也让我看到了自身的不足和未来的成长空间。在复杂图纸识读、精密工艺分析、独立问题解决和项目管理能力等方面，我与经验丰富的工程师之间还存在明显差距。但我并不气馁，因为差距就是成长的方向，明确了不足才能有针对性地提升。')
body('此次实习也让我对机械工程专业和制造业发展前景有了更加坚定的信心。随着中国制造业向智能化、精密化、绿色化方向转型升级，企业对工程技术人才的需求将持续增长。工程助理作为衔接技术与生产的核心岗位，是进入工程技术领域最直接有效的通道。')
body('最后，衷心感谢东莞浩镒晟五金电子制品有限公司为我提供的宝贵实习平台，感谢公司领导和工程部主管在工作中的信任与指导，感谢每一位同事在工作和生活上的关心与帮助。我将带着这段实习赋予我的专业知识、实操技能和职业信念，满怀信心地走向未来的工程技术职业道路。')

# ===== SAVE =====
import os
outpath = r'C:/Users/lb/Desktop/岗位实习总结_带图定稿.docx'
doc.save(outpath)

total = sum(len(p.text.replace(' ','').replace('\n','')) for p in doc.paragraphs)
# Count images
img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        if r._r.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            img_count += 1
print(f'Saved: {len(doc.paragraphs)}p ~{total}c {img_count} imgs')
print(f'Path: {outpath}')
