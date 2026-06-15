#!/usr/bin/env python
"""Fix all tables in course design report with real OCR data."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

path = r'C:/Users/lb/Desktop/电力系统工程综合课程设计_新居.docx'
doc = Document(path)
F = '仿宋_GB2312'

def cp(p):
    for r in list(p.runs): r._r.getparent().remove(r._r)
def ar(p, text):
    r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:eastAsia'), F); rF.set(qn('w:ascii'), F); rF.set(qn('w:hAnsi'), F)
    rPr.append(rF); s = OxmlElement('w:sz'); s.set(qn('w:val'), '24'); rPr.append(s)
    r.append(rPr); t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); p._element.append(r)

def set_cell(table, row, col, value):
    for p in table.cell(row, col).paragraphs:
        for r in p.runs: r.text = value

# ===== TABLE 1 (load) 18 rows =====
data1 = [
    ['走廊','照明灯','20','0.165'],
    ['主卧室','电视机','200','1.136'],
    ['主卧室','空调','2100','16.970'],
    ['主次卧室','照明及插座','600','2.727'],
    ['次卧室','空调','1800','14.545'],
    ['书房','电脑及照明','400','2.273'],
    ['书房','空调','1800','14.545'],
    ['厨房','抽油烟机','280','2.263'],
    ['厨房','微波炉','800','3.636'],
    ['厨房','电饭煲','900','4.091'],
    ['厨房','电磁炉','2000','9.091'],
    ['厨房','电水壶','1500','6.819'],
    ['厨房','冰箱','110','0.889'],
    ['卫生间','浴霸','1000','4.545'],
    ['卫生间','吹风机','1000','8.081'],
    ['客厅','电视机及插座','400','2.273'],
    ['客厅','空调','2400','19.394'],
    ['阳台','电动车+洗衣机','600','3.611'],
]
for i, row in enumerate(data1):
    for j in range(min(4, len(row))):
        set_cell(doc.tables[1], i+1, j, row[j])
print('1. Load table')

# ===== TABLE 2 (area) 8 rows =====
data2 = [
    ('走廊','S=(1.5-0.24)*(2.2-0.24)=2.47'),
    ('主卧室','S=(3.4-0.24)*(3.6-0.24)=10.62'),
    ('次卧室','S=(3.0-0.24)*(3.3-0.24)=8.45'),
    ('书房','S=(3.2-0.24)*(3.3-0.24)=9.06'),
    ('厨房','S=(1.85-0.24)*(2.8-0.24)=4.12'),
    ('卫生间','S=(1.7-0.24)*(2.0-0.24)=2.57'),
    ('客厅','S=(3.9-0.24)*(4.5-0.24)=15.59'),
    ('阳台','S=(10.5-0.24)*(1.5-0.24)=12.93'),
]
for i, (n, a) in enumerate(data2):
    set_cell(doc.tables[2], i+1, 0, n)
    set_cell(doc.tables[2], i+1, 1, a)
print('2. Area table')

# ===== TABLE 4 (lighting) 13 rows =====
data4 = [
    ['走廊','照明灯','2.47','20','0.859','5W节能灯'],
    ['主卧室','照明灯','10.62','50','9.235','9W节能灯'],
    ['主卧室','书桌灯','4','100','6.957','7W节能灯'],
    ['次卧室','照明灯','8.45','50','7.348','7W节能灯'],
    ['次卧室','书桌灯','4','100','6.957','7W节能灯'],
    ['书房','照明灯','9.06','50','7.878','7W节能灯'],
    ['书房','书桌灯','4','100','6.957','7W节能灯'],
    ['厨房','照明灯','4.12','50','16.480','18W白炽灯'],
    ['卫生间','照明灯','2.57','15','0.670','5W节能灯'],
    ['卫生间','镜子灯','4','50','3.478','5W节能灯'],
    ['客厅','照明灯','15.59','50','13.557','14W节能灯'],
    ['客厅','看书灯','4','100','6.957','7W节能灯'],
    ['阳台','照明灯','12.93','20','4.497','5W节能灯'],
]
for i, row in enumerate(data4):
    for j, val in enumerate(row):
        set_cell(doc.tables[4], i+1, j, val)
print('3. Lighting table')

# Fix total lighting para
cp(doc.paragraphs[137])
ar(doc.paragraphs[137], 'P总=0.859+9.235+6.957+7.348+6.957+7.878+6.957+16.480+0.670+3.478+13.557+6.957+4.497=91.873W')
cp(doc.paragraphs[138]); ar(doc.paragraphs[138], '')

# ===== TABLE 5 (breaker) 12 rows =====
data5 = [
    ['L1','主卧空调','16.970','20A/1P','4mm2'],
    ['L2','次卧空调','14.545','20A/1P','4mm2'],
    ['L3','书房空调','14.545','20A/1P','4mm2'],
    ['L4','客厅空调','19.394','25A/1P','4mm2'],
    ['L5','电动车充电','1.591','10A/1P','1.5mm2'],
    ['L6','卫生间浴霸','4.545','10A/1P+漏保','1.5mm2'],
    ['L7','卫生间吹风机','8.081','10A/1P','1.5mm2'],
    ['L8','厨房大功率一','15.910','20A/1P+漏保','4mm2'],
    ['L9','厨房大功率二','3.636','10A/1P','1.5mm2'],
    ['L10','厨房普通+照明','7.243','10A/1P+漏保','2.5mm2'],
    ['L11','卧室书房小功率','19.148','25A/1P','4mm2'],
    ['L12','公共区域+阳台','13.119','16A/1P','2.5mm2'],
]
for i, row in enumerate(data5):
    for j, val in enumerate(row):
        set_cell(doc.tables[5], i+1, j, val)
print('4. Breaker table')

# ===== FIX I_total =====
vals = [0.165,1.136,16.970,2.727,14.545,2.273,14.545,2.263,3.636,4.091,9.091,6.819,0.889,4.545,8.081,2.273,19.394,3.611]
rem = sum(vals)
it = 19.394 + 0.5 * rem
print(f'I_total={it:.3f}A')

cp(doc.paragraphs[122]); ar(doc.paragraphs[122], f'I总=19.394+0.5*({rem:.3f})')
cp(doc.paragraphs[123]); ar(doc.paragraphs[123], f'  =19.394+{0.5*rem:.3f}')
cp(doc.paragraphs[124]); ar(doc.paragraphs[124], f'  ={it:.3f}A')
cp(doc.paragraphs[125]); ar(doc.paragraphs[125], '')

# Fix body references
for i in range(200, 212):
    try:
        for r in doc.paragraphs[i].runs:
            t = r.text or ''
            t = t.replace('69.35', f'{it:.1f}').replace('69.349', f'{it:.3f}')
            t = t.replace('107.83', '91.87').replace('107.826', '91.873')
            if t != r.text: r.text = t
    except: pass

# Fix circuit description
cp(doc.paragraphs[140])
ar(doc.paragraphs[140], '该三口人家庭户型为L型布局（左侧上部凹入），含客厅、主次卧室、书房、厨房、卫生间及阳台共7区域，建筑面积约78m2。大功率电器包括4台空调（客厅2400W+主卧2100W+次卧1800W+书房1800W）、厨房电磁炉2000W及电水壶1500W、卫生间浴霸1000W及吹风机1000W，均单独回路，计12条。')

cp(doc.paragraphs[141])
ar(doc.paragraphs[141], '12条回路分配如下：')
loops = [
    '（1）主卧空调（L1）；','（2）次卧空调（L2）；','（3）书房空调（L3）；',
    '（4）客厅空调（L4）；','（5）电动车充电（L5）；',
    '（6）卫生间浴霸（L6）；','（7）卫生间吹风机（L7）；',
    '（8）厨房电磁炉+电水壶（L8）；','（9）厨房微波炉（L9）；',
    '（10）厨房电饭煲+抽烟机+照明（L10）；',
    '（11）主次卧+书房照明及普通插座（L11）；',
    '（12）走廊+客厅+卫+阳台照明及普通插座（L12）。',
]
for j, txt in enumerate(loops):
    cp(doc.paragraphs[142+j]); ar(doc.paragraphs[142+j], txt)

outpath = r'C:/Users/lb/Desktop/电力系统工程课程设计_修正版.docx'
doc.save(outpath)
print(f'\nSaved: {outpath}')
print('ALL TABLES AND DATA CORRECTED.')
