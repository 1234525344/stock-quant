# -*- coding: utf-8 -*-
"""
将生成的实验截图插入DOCX报告
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

IMG_DIR = r"C:\Users\lb\stock-quant\report_images"

# 每个实验需要插入的图片
exp_images = {
    0: [  # 实验1
        ("exp1_wiring.png", "实验1 PLC接线图"),
        ("exp1_ladder.png", "实验1 梯形图程序截图"),
        ("exp1_debug.png", "实验1 TIA Portal在线监控截图"),
        ("exp1_real.png", "实验1 实物接线图"),
    ],
    1: [  # 实验2
        ("exp2_wiring.png", "实验2 PLC接线图"),
        ("exp2_ladder.png", "实验2 梯形图程序截图"),
        ("exp2_debug.png", "实验2 TIA Portal在线监控截图"),
    ],
    2: [  # 实验3
        ("exp3_wiring.png", "实验3 PLC接线图"),
        ("exp3_ladder.png", "实验3 梯形图程序截图"),
        ("exp3_debug.png", "实验3 TIA Portal在线监控截图"),
    ],
    3: [  # 实验4
        ("exp4_wiring.png", "实验4 PLC接线图"),
        ("exp4_ladder.png", "实验4 梯形图程序截图"),
        ("exp4_debug.png", "实验4 TIA Portal在线监控截图"),
    ],
}


def insert_images(template_path, output_path):
    """在报告的每个实验数据表格中插入图片"""
    doc = Document(template_path)

    # 数据表格索引：2, 4, 6, 8
    data_table_indices = [2, 4, 6, 8]

    for exp_idx, table_idx in enumerate(data_table_indices):
        if exp_idx >= len(exp_images):
            break
        if table_idx >= len(doc.tables):
            break

        table = doc.tables[table_idx]
        cell_data = table.cell(0, 1)  # "实验数据与结论"的内容单元格

        # 在单元格末尾添加图片
        for img_file, caption in exp_images[exp_idx]:
            img_path = os.path.join(IMG_DIR, img_file)
            if not os.path.exists(img_path):
                print(f"  ⚠️ 图片不存在: {img_path}")
                continue

            # 添加标题
            p = cell_data.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"\n{caption}")
            run.font.size = Pt(10)
            run.bold = True

            # 添加图片
            p = cell_data.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(5.5))

            print(f"  ✅ 插入: {caption}")

    doc.save(output_path)
    print(f"\n✅ 报告已保存: {output_path}")


if __name__ == "__main__":
    template = r"C:\Users\lb\Documents\PLC实验报告册-v3.docx"
    output = r"C:\Users\lb\Documents\PLC实验报告册-v3完整版.docx"

    insert_images(template, output)
